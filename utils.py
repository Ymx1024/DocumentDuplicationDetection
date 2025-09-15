# 提前在子进程中抑制初始化日志，避免重复噪声
try:
    import os as _early_os
    import multiprocessing as _early_mp
    if _early_os.getenv('SUPPRESS_INIT_LOGS', '0') == '1' and getattr(_early_mp.current_process(), 'name', 'MainProcess') != 'MainProcess':
        def print(*args, **kwargs):  # type: ignore[override]
            pass
except Exception:
    pass

print("🔧 [DEBUG] 开始导入utils.py模块...")

import os
import re
import sys
import difflib
print("🔧 [DEBUG] 基础模块导入完成 (os, re, sys, difflib)")

try:
    import diff_match_patch as dmp_module
    _DMP_AVAILABLE = True
    print("🔧 [DEBUG] diff-match-patch 导入成功")
except ImportError:
    _DMP_AVAILABLE = False
    print("🔧 [DEBUG] diff-match-patch 不可用，将使用 difflib 作为备用方案")

print("🔧 [DEBUG] 准备导入sentence_transformers (这可能需要一些时间)...")
try:
    from sentence_transformers import SentenceTransformer, util
    print("🔧 [DEBUG] sentence_transformers 导入成功")
except ImportError as e:
    print(f"🔧 [DEBUG] sentence_transformers 导入失败: {e}")
    raise

print("🔧 [DEBUG] 开始导入torch...")
import torch
print("🔧 [DEBUG] torch 导入成功")

print("🔧 [DEBUG] 开始导入nltk...")
try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    print("🔧 [DEBUG] nltk 导入成功")
except ImportError as e:
    print(f"🔧 [DEBUG] nltk 导入失败: {e}")

print("🔧 [DEBUG] 开始导入其他依赖...")
import yaml
import pandas as pd
from collections import Counter
import uuid
from docx import Document
import textract
import json
import shutil
import multiprocessing
import threading
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor
import functools
import warnings
from multiprocessing import Process, Queue
import time
print("🔧 [DEBUG] 所有导入完成")

# 抑制multiprocessing相关的警告
warnings.filterwarnings("ignore", category=UserWarning, module="multiprocessing")
warnings.filterwarnings("ignore", message="resource_tracker: process died unexpectedly")
warnings.filterwarnings("ignore", message="Some resources might leak")

# 彻底规避Python3.8+在joblib/loky下偶发的resource_tracker KeyError噪声
try:
    import multiprocessing.resource_tracker as _rt  # type: ignore
    # 仅打一次补丁
    if not getattr(_rt, "_patched_ignore_loky_semlock", False):
        _orig_register = _rt.register
        _orig_unregister = _rt.unregister

        def _safe_register(name, rtype):
            # loky自行管理semlock，避免重复注册
            if rtype == 'semlock':
                return
            try:
                return _orig_register(name, rtype)
            except Exception:
                # 避免任何注册阶段的异常影响主流程
                return

        def _safe_unregister(name, rtype):
            try:
                return _orig_unregister(name, rtype)
            except Exception:
                # 关键点：忽略退出阶段的KeyError等异常
                return

        _rt.register = _safe_register  # type: ignore
        _rt.unregister = _safe_unregister  # type: ignore
        _rt._patched_ignore_loky_semlock = True  # type: ignore
except Exception:
    # 如果打补丁失败，不影响主流程
    pass

# 设置环境变量来减少警告
os.environ['PYTHONWARNINGS'] = 'ignore::UserWarning'

# 更严格的主进程检测：只有在真正的主进程启动时才显示初始化日志
def is_main_initialization_process():
    """检测是否为主初始化进程，避免在多进程任务中重复显示日志"""
    print(f"🔧 [DEBUG] 检查主进程状态: SUPPRESS_INIT_LOGS={os.getenv('SUPPRESS_INIT_LOGS', '0')}")
    
    # 如果已经设置了抑制标志，直接返回False
    if _should_suppress_logs:
        print("🔧 [DEBUG] 抑制标志已设置，跳过初始化")
        return False
    
    # 检查是否在multiprocessing子进程中
    if hasattr(multiprocessing.current_process(), 'name'):
        process_name = multiprocessing.current_process().name
        print(f"🔧 [DEBUG] 当前进程名: {process_name}")
        if process_name != 'MainProcess':
            print("🔧 [DEBUG] 非主进程，跳过初始化")
            return False
    
    # 检查是否在Flask应用启动过程中
    is_flask_main = __name__ == '__main__' or os.getenv('FLASK_ENV') or 'werkzeug' in sys.modules
    print(f"🔧 [DEBUG] Flask主进程检查: __name__={__name__}, FLASK_ENV={os.getenv('FLASK_ENV')}, werkzeug in sys.modules={'werkzeug' in sys.modules}")
    print(f"🔧 [DEBUG] 最终判断结果: {is_flask_main}")
    
    return is_flask_main

# 在模块导入时就检查是否应该抑制日志，避免在子进程中重复显示
_should_suppress_logs = os.getenv('SUPPRESS_INIT_LOGS', '0') == '1'

# 如果设置了抑制标志，直接跳过所有初始化日志
if _should_suppress_logs:
    # 静默模式：不打印任何初始化日志
    def silent_print(*args, **kwargs):
        pass
    # 临时替换print函数
    _original_print = print
    print = silent_print
import psutil
import time

# 设置多进程启动方法为spawn，解决CUDA fork问题
try:
    multiprocessing.set_start_method('spawn', force=True)
except RuntimeError:
    # 如果已经设置过，忽略错误
    pass

# 推荐多线程设置（避免仅单核）
try:
    CPU_CORES = multiprocessing.cpu_count()
    WORKER_THREADS = min(CPU_CORES, 16)  # 增加工作线程数
    
    # 强制设置所有相关的线程数环境变量
    os.environ['OMP_NUM_THREADS'] = str(WORKER_THREADS)
    os.environ['MKL_NUM_THREADS'] = str(WORKER_THREADS)
    os.environ['NUMEXPR_NUM_THREADS'] = str(WORKER_THREADS)
    os.environ['OPENBLAS_NUM_THREADS'] = str(WORKER_THREADS)
    os.environ['VECLIB_MAXIMUM_THREADS'] = str(WORKER_THREADS)
    os.environ['NUMBA_NUM_THREADS'] = str(WORKER_THREADS)
    
    # 设置PyTorch线程数
    torch.set_num_threads(WORKER_THREADS)
    torch.set_num_interop_threads(WORKER_THREADS)
    
    # 只在主进程中打印多线程配置信息
    if is_main_initialization_process():
        print(f"检测到CPU核心数: {CPU_CORES}")
        print(f"使用工作线程数: {WORKER_THREADS}")
        print(f"环境变量设置: OMP_NUM_THREADS={os.environ.get('OMP_NUM_THREADS')}")
except Exception as _e:
    print(f"多线程配置失败: {_e}")

# 允许在无网络时降级NLTK与模型依赖
print("🔧 [DEBUG] 开始检查NLTK资源...")
try:
    # 首先检查是否已经有NLTK数据
    try:
        from nltk.tokenize import sent_tokenize, word_tokenize
        # 快速测试是否可用
        sent_tokenize("Test sentence.")
        word_tokenize("Test sentence.")
        print("🔧 [DEBUG] NLTK资源已存在，跳过下载")
        _nltk_ok = True
    except LookupError:
        # 数据不存在，需要下载
        print("🔧 [DEBUG] NLTK资源不存在，开始下载...")
        
        # 使用超时机制下载NLTK数据
        import signal
        
        def timeout_handler(signum, frame):
            raise TimeoutError("NLTK下载超时")
        
        try:
            # 设置30秒超时
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(30)
            
            print("🔧 [DEBUG] 下载punkt数据包...")
            nltk.download('punkt', quiet=True)
            print("🔧 [DEBUG] 下载averaged_perceptron_tagger数据包...")
            nltk.download('averaged_perceptron_tagger', quiet=True)
            
            signal.alarm(0)  # 取消超时
            print("🔧 [DEBUG] NLTK资源下载完成")
            _nltk_ok = True
            
        except (TimeoutError, Exception) as download_error:
            signal.alarm(0)  # 确保取消超时
            print(f"🔧 [DEBUG] NLTK下载失败或超时: {download_error}")
            print("🔧 [DEBUG] 将使用简易分词作为后备方案")
            _nltk_ok = False
except Exception as _e:
    print(f"🔧 [DEBUG] NLTK资源下载失败，使用简易分词降级: {_e}")
    _nltk_ok = False

print("🔧 [DEBUG] NLTK检查完成")

# 简易分句/分词降级实现
print("🔧 [DEBUG] 导入re模块...")
import re as _re
print("🔧 [DEBUG] 定义简易分词函数...")

def _simple_sent_tokenize(text: str):
    return [s for s in _re.split(r"(?<=[。！？.!?])\s+", text) if s.strip()]

print("🔧 [DEBUG] 定义_simple_sent_tokenize完成")

def _simple_word_tokenize(text: str):
    return [w for w in _re.findall(r"[\w\-]+", text.lower())]

print("🔧 [DEBUG] 定义_simple_word_tokenize完成")

# 根据可用性选择分词器
print("🔧 [DEBUG] 配置分词器...")
_sent_tokenize = sent_tokenize if _nltk_ok else _simple_sent_tokenize
_word_tokenize = word_tokenize if _nltk_ok else _simple_word_tokenize
print("🔧 [DEBUG] 分词器配置完成")

print("🔧 [DEBUG] 定义monitor_cpu_usage函数...")
def monitor_cpu_usage():
    """监控CPU使用情况"""
    try:
        cpu_percent = psutil.cpu_percent(interval=1, percpu=True)
        avg_cpu = sum(cpu_percent) / len(cpu_percent)
        active_cores = sum(1 for cpu in cpu_percent if cpu > 10)
        print(f"CPU使用率: 平均 {avg_cpu:.1f}%, 活跃核心数: {active_cores}/{len(cpu_percent)}")
        
        # 显示每个核心的使用情况
        if len(cpu_percent) <= 16:  # 只显示前16个核心
            core_usage = [f"{i}:{cpu:.1f}%" for i, cpu in enumerate(cpu_percent[:16])]
            print(f"核心使用详情: {', '.join(core_usage)}")
        
        return avg_cpu, active_cores
    except Exception as e:
        print(f"CPU监控失败: {e}")
        return 0, 0

def monitor_gpu_memory():
    """监控GPU内存使用情况"""
    try:
        if torch.cuda.is_available():
            gpu_memory_used = torch.cuda.memory_allocated() / 1024**3  # GB
            gpu_memory_total = torch.cuda.get_device_properties(0).total_memory / 1024**3  # GB
            gpu_memory_free = gpu_memory_total - gpu_memory_used
            print(f"GPU内存: 已用 {gpu_memory_used:.2f}GB, 总计 {gpu_memory_total:.2f}GB, 可用 {gpu_memory_free:.2f}GB")
            return gpu_memory_used, gpu_memory_total, gpu_memory_free
        else:
            print("GPU不可用")
            return 0, 0, 0
    except Exception as e:
        print(f"GPU监控失败: {e}")
        return 0, 0, 0

def cleanup_gpu_memory():
    """清理GPU内存"""
    try:
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.synchronize()
            print("GPU内存已清理")
    except Exception as e:
        print(f"GPU内存清理失败: {e}")

def extract_text_parallel(file_paths, max_workers=None):
    """并行提取多个文件的文本"""
    if max_workers is None:
        max_workers = min(multiprocessing.cpu_count(), 8)
    
    def extract_single_file(file_path):
        try:
            lines, html, temp_path = extract_formatted_text(file_path, is_storage_file=True)
            return file_path, lines, html, temp_path, None
        except Exception as e:
            return file_path, None, None, None, str(e)
    
    with ProcessPoolExecutor(max_workers=max_workers) as executor:
        results = list(executor.map(extract_single_file, file_paths))
    
    return results

# 配置
UPLOAD_FOLDER = 'Uploads'
STORAGE_FOLDER = 'storage'
SIMILARITIES_FILE = 'similarities.json'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'txt', 'yaml', 'yml', 'xlsx', 'xls'}
MAX_SENTENCES_PER_SEGMENT = 20
SIMILARITY_THRESHOLD = 0.95

# 全局任务取消状态存储
CANCELLED_TASKS = set()
_CANCELLED_TASKS_LOCK = threading.Lock()

# 任务超时机制
TASK_TIMEOUTS = {}  # {task_id: start_time}
_TASK_TIMEOUTS_LOCK = threading.Lock()
TASK_TIMEOUT_SECONDS = 300  # 5分钟超时，适合大文件处理任务

# 多用户环境管理
USER_SESSIONS = {}  # {user_id: {tasks: set, last_activity: timestamp, ip: str}}
_USER_SESSIONS_LOCK = threading.Lock()
TASK_TO_USER = {}  # {task_id: user_id}
_TASK_TO_USER_LOCK = threading.Lock()

# 系统资源监控
SYSTEM_RESOURCES = {
    'active_tasks': 0,
    'active_users': 0,
    'cpu_usage': 0.0,
    'memory_usage': 0.0,
    'gpu_memory_usage': 0.0
}
_SYSTEM_RESOURCES_LOCK = threading.Lock()

# 进程管理机制
ACTIVE_PROCESSES = {}  # {task_id: [process_objects]}
_ACTIVE_PROCESSES_LOCK = threading.Lock()

# 超时检查线程
_timeout_checker_thread = None
_timeout_checker_running = False
_timeout_checker_lock = threading.Lock()

# 用户清理线程
_user_cleanup_thread = None
_user_cleanup_running = False
_user_cleanup_lock = threading.Lock()

def is_task_cancelled(task_id):
    """检查任务是否已被取消"""
    if not task_id:
        return False
    with _CANCELLED_TASKS_LOCK:
        return task_id in CANCELLED_TASKS

def mark_task_cancelled(task_id):
    """标记任务为已取消"""
    if task_id:
        with _CANCELLED_TASKS_LOCK:
            CANCELLED_TASKS.add(task_id)
            print(f"任务 {task_id} 已标记为取消")
            # 强制终止所有相关进程
            terminate_task_processes(task_id)

def clear_cancelled_task(task_id):
    """清除任务的取消状态"""
    if task_id:
        with _CANCELLED_TASKS_LOCK:
            CANCELLED_TASKS.discard(task_id)

def get_user_id_from_request(request):
    """从请求中获取用户ID（基于IP地址和User-Agent）"""
    import hashlib
    ip = request.remote_addr or 'unknown'
    user_agent = request.headers.get('User-Agent', 'unknown')
    user_string = f"{ip}_{user_agent}"
    return hashlib.md5(user_string.encode()).hexdigest()[:16]

def register_user_session(user_id, request):
    """注册用户会话"""
    with _USER_SESSIONS_LOCK:
        if user_id not in USER_SESSIONS:
            USER_SESSIONS[user_id] = {
                'tasks': set(),
                'last_activity': time.time(),
                'ip': request.remote_addr or 'unknown',
                'user_agent': request.headers.get('User-Agent', 'unknown')
            }
        else:
            USER_SESSIONS[user_id]['last_activity'] = time.time()
            USER_SESSIONS[user_id]['ip'] = request.remote_addr or 'unknown'
        
        # 更新系统资源统计
        with _SYSTEM_RESOURCES_LOCK:
            SYSTEM_RESOURCES['active_users'] = len(USER_SESSIONS)
    
    print(f"用户会话已注册: {user_id} (IP: {request.remote_addr})")

def register_task_to_user(task_id, user_id):
    """将任务注册到用户"""
    with _TASK_TO_USER_LOCK:
        TASK_TO_USER[task_id] = user_id
    
    with _USER_SESSIONS_LOCK:
        if user_id in USER_SESSIONS:
            USER_SESSIONS[user_id]['tasks'].add(task_id)
            USER_SESSIONS[user_id]['last_activity'] = time.time()
    
    # 更新系统资源统计
    with _SYSTEM_RESOURCES_LOCK:
        SYSTEM_RESOURCES['active_tasks'] = len(TASK_TIMEOUTS)
    
    print(f"任务 {task_id} 已注册到用户 {user_id}")

def cancel_user_tasks(user_id):
    """取消指定用户的所有任务"""
    cancelled_count = 0
    
    with _USER_SESSIONS_LOCK:
        if user_id not in USER_SESSIONS:
            return 0
        
        user_tasks = list(USER_SESSIONS[user_id]['tasks'])
    
    print(f"用户 {user_id} 页面刷新：开始取消 {len(user_tasks)} 个任务")
    
    # 保存任务列表用于进程清理
    tasks_for_cleanup = user_tasks.copy()
    
    for task_id in user_tasks:
        print(f"用户 {user_id} 页面刷新：取消任务 {task_id}")
        mark_task_cancelled(task_id)
        cancelled_count += 1
        
        # 清理任务计时
        with _TASK_TIMEOUTS_LOCK:
            TASK_TIMEOUTS.pop(task_id, None)
        
        # 清理任务到用户的映射
        with _TASK_TO_USER_LOCK:
            TASK_TO_USER.pop(task_id, None)
        
        # 终止任务相关进程
        terminate_task_processes(task_id)
        clear_task_processes(task_id)
    
    # 清理用户会话中的任务
    with _USER_SESSIONS_LOCK:
        if user_id in USER_SESSIONS:
            USER_SESSIONS[user_id]['tasks'].clear()
            USER_SESSIONS[user_id]['last_activity'] = time.time()
    
    if cancelled_count > 0:
        print(f"用户 {user_id} 页面刷新：已取消 {cancelled_count} 个任务")
        
        # 更新系统资源统计
        with _SYSTEM_RESOURCES_LOCK:
            SYSTEM_RESOURCES['active_tasks'] = len(TASK_TIMEOUTS)
        
        # 使用保存的任务列表进行进程清理
        killed_count = force_cleanup_user_processes_with_tasks(user_id, tasks_for_cleanup)
        print(f"用户 {user_id} 进程清理：已终止 {killed_count} 个进程")
        
        # 额外等待确保进程被清理
        time.sleep(0.5)
        
        # 再次强制清理multiprocessing进程
        kill_all_multiprocessing_processes()
        
        # 强制清理所有pt_main_thread进程
        force_kill_pt_main_thread_processes_with_tasks(user_id, tasks_for_cleanup)
        
        # 额外清理：杀死所有pt_main_thread进程（除了主服务进程）
        force_kill_all_pt_main_thread_processes_safe()
        
        print(f"用户 {user_id} 页面刷新：已完成进程清理")
    
    return cancelled_count

def cancel_all_running_tasks():
    """取消所有正在运行的任务（用于页面刷新检测）"""
    cancelled_count = 0
    with _TASK_TIMEOUTS_LOCK:
        # 获取所有正在运行的任务ID
        running_tasks = list(TASK_TIMEOUTS.keys())
        
    for task_id in running_tasks:
        print(f"页面刷新检测：取消正在运行的任务 {task_id}")
        mark_task_cancelled(task_id)
        cancelled_count += 1
        
        # 清理任务计时
        with _TASK_TIMEOUTS_LOCK:
            TASK_TIMEOUTS.pop(task_id, None)
        
        # 终止任务相关进程
        terminate_task_processes(task_id)
        clear_task_processes(task_id)
    
    if cancelled_count > 0:
        print(f"页面刷新检测：已取消 {cancelled_count} 个正在运行的任务")
        
        # 立即强制清理所有multiprocessing进程
        print("页面刷新检测：开始强制清理multiprocessing进程...")
        kill_all_multiprocessing_processes()
        
        # 强制清理所有Python进程
        print("页面刷新检测：开始强制清理Python进程...")
        force_kill_all_python_processes()
        
        # 额外等待一小段时间确保进程被清理
        time.sleep(0.5)
        
        print("页面刷新检测：进程清理完成")
    
    return cancelled_count

def force_cleanup_user_processes(user_id):
    """强制清理指定用户的进程"""
    try:
        import psutil
        import signal
        current_pid = os.getpid()
        killed_count = 0
        
        # 获取该用户的所有任务
        with _USER_SESSIONS_LOCK:
            if user_id not in USER_SESSIONS:
                return 0
            user_tasks = list(USER_SESSIONS[user_id]['tasks'])
        
        print(f"用户 {user_id} 进程清理：开始清理任务 {user_tasks}")
        
        # 第一轮：查找并终止该用户的直接子进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['ppid'] == current_pid and proc.info['name'] == 'python':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    # 检查是否是multiprocessing相关进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 进程清理：发现相关进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        proc.terminate()
                        killed_count += 1
                        # 等待进程终止
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        # 第二轮：查找所有包含任务ID的进程（包括孤儿进程）
        for task_id in user_tasks:
            for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
                try:
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    if task_id in cmdline and proc.info['name'] == 'python':
                        print(f"用户 {user_id} 进程清理：发现任务 {task_id} 的进程 {proc.info['pid']} (PPID: {proc.info['ppid']})")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"进程 {proc.info['pid']} 无法被杀死")
                except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                    pass
        
        # 第三轮：强制清理所有pt_main_thread进程（这些通常是PyTorch相关进程）
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    # 检查是否是我们的进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 进程清理：发现pt_main_thread进程 {proc.info['pid']}")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"用户 {user_id} 进程清理：已终止 {killed_count} 个进程")
        
        return killed_count
    except Exception as e:
        print(f"用户 {user_id} 进程清理时出错: {e}")
        return 0

def force_cleanup_user_processes_with_tasks(user_id, user_tasks):
    """使用指定任务列表强制清理用户进程"""
    try:
        import psutil
        import signal
        current_pid = os.getpid()
        killed_count = 0
        
        print(f"用户 {user_id} 进程清理：开始清理任务 {user_tasks}")
        
        # 第一轮：查找并终止该用户的直接子进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['ppid'] == current_pid and proc.info['name'] == 'python':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    # 检查是否是multiprocessing相关进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 进程清理：发现相关进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        proc.terminate()
                        killed_count += 1
                        # 等待进程终止
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        # 第二轮：查找所有包含任务ID的进程（包括孤儿进程）
        for task_id in user_tasks:
            for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
                try:
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    if task_id in cmdline and proc.info['name'] == 'python':
                        print(f"用户 {user_id} 进程清理：发现任务 {task_id} 的进程 {proc.info['pid']} (PPID: {proc.info['ppid']})")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"进程 {proc.info['pid']} 无法被杀死")
                except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                    pass
        
        # 第三轮：强制清理所有pt_main_thread进程（这些通常是PyTorch相关进程）
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    # 检查是否是我们的进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 进程清理：发现pt_main_thread进程 {proc.info['pid']}")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"用户 {user_id} 进程清理：已终止 {killed_count} 个进程")
        
        return killed_count
    except Exception as e:
        print(f"用户 {user_id} 进程清理时出错: {e}")
        return 0

def force_kill_pt_main_thread_processes(user_id):
    """强制杀死所有pt_main_thread进程"""
    try:
        import psutil
        killed_count = 0
        
        # 获取该用户的所有任务
        with _USER_SESSIONS_LOCK:
            if user_id not in USER_SESSIONS:
                return 0
            user_tasks = list(USER_SESSIONS[user_id]['tasks'])
        
        print(f"用户 {user_id} 强制清理pt_main_thread进程：开始清理任务 {user_tasks}")
        
        # 查找所有pt_main_thread进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    # 检查是否是我们的进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 强制清理：发现pt_main_thread进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"pt_main_thread进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"用户 {user_id} 强制清理pt_main_thread进程：已终止 {killed_count} 个进程")
        
        return killed_count
    except Exception as e:
        print(f"用户 {user_id} 强制清理pt_main_thread进程时出错: {e}")
        return 0

def force_kill_pt_main_thread_processes_with_tasks(user_id, user_tasks):
    """使用指定任务列表强制杀死所有pt_main_thread进程"""
    try:
        import psutil
        killed_count = 0
        current_pid = os.getpid()
        
        print(f"用户 {user_id} 强制清理pt_main_thread进程：开始清理任务 {user_tasks}")
        
        # 查找所有pt_main_thread进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    
                    # 跳过主服务进程
                    if proc.info['pid'] == current_pid:
                        continue
                    
                    # 跳过包含app.py或start_app.py的进程（主服务进程）
                    if 'app.py' in cmdline or 'start_app.py' in cmdline:
                        continue
                    
                    # 只处理子进程
                    if proc.info['ppid'] != current_pid:
                        continue
                    
                    # 检查是否是我们的进程
                    if any(task_id in cmdline for task_id in user_tasks):
                        print(f"用户 {user_id} 强制清理：发现pt_main_thread进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"pt_main_thread进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"用户 {user_id} 强制清理pt_main_thread进程：已终止 {killed_count} 个进程")
        
        return killed_count
    except Exception as e:
        print(f"用户 {user_id} 强制清理pt_main_thread进程时出错: {e}")
        return 0

def force_kill_all_pt_main_thread_processes():
    """强制杀死所有pt_main_thread进程（激进方法）"""
    try:
        import psutil
        killed_count = 0
        current_pid = os.getpid()
        
        print("强制清理：开始杀死所有pt_main_thread进程")
        
        # 查找所有pt_main_thread进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    
                    # 跳过主服务进程
                    if proc.info['pid'] == current_pid:
                        print(f"强制清理：跳过主服务进程 {proc.info['pid']}")
                        continue
                    
                    # 跳过包含app.py或start_app.py的进程（主服务进程）
                    if 'app.py' in cmdline or 'start_app.py' in cmdline:
                        print(f"强制清理：跳过主服务进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        continue
                    
                    # 只杀死子进程
                    if proc.info['ppid'] == current_pid:
                        print(f"强制清理：发现子pt_main_thread进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        proc.terminate()
                        killed_count += 1
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"pt_main_thread进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            try:
                                proc.wait(timeout=1)
                            except psutil.TimeoutExpired:
                                print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
                    else:
                        print(f"强制清理：跳过非子进程 {proc.info['pid']} (PPID: {proc.info['ppid']})")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"强制清理：已终止 {killed_count} 个pt_main_thread进程")
        
        return killed_count
    except Exception as e:
        print(f"强制清理pt_main_thread进程时出错: {e}")
        return 0

def cleanup_task_specific_processes(task_id):
    """温和地清理特定任务相关的进程，不影响系统进程"""
    try:
        import psutil
        killed_count = 0
        current_pid = os.getpid()
        
        print(f"温和清理：开始清理任务 {task_id} 的相关进程")
        
        # 只查找明确包含任务ID的进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['pid'] == current_pid:
                    continue
                    
                cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                
                # 只清理明确包含任务ID的进程
                if task_id in cmdline and proc.info['name'] == 'python':
                    # 额外检查：跳过重要的系统进程
                    if any(system_keyword in cmdline for system_keyword in [
                        'resource_tracker', 'forkserver', 'semlock_tracker', 'app.py', 'start_app.py'
                    ]):
                        print(f"温和清理：跳过系统进程 {proc.info['pid']}: {cmdline[:100]}...")
                        continue
                    
                    print(f"温和清理：发现任务进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                    proc.terminate()
                    killed_count += 1
                    try:
                        proc.wait(timeout=2)
                    except psutil.TimeoutExpired:
                        print(f"温和清理：进程 {proc.info['pid']} 未响应，跳过强制杀死")
                        # 温和清理模式下不使用 kill()
                        
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"温和清理：已终止 {killed_count} 个任务相关进程")
        else:
            print("温和清理：未找到需要清理的任务进程")
        
        return killed_count
    except Exception as e:
        print(f"温和清理任务进程时出错: {e}")
        return 0

def force_kill_all_pt_main_thread_processes_safe():
    """安全地杀死所有pt_main_thread进程（除了主服务进程）"""
    try:
        import psutil
        killed_count = 0
        current_pid = os.getpid()
        
        print("安全清理：开始杀死所有pt_main_thread进程（除了主服务进程）")
        
        # 查找所有pt_main_thread进程
        for proc in psutil.process_iter(['pid', 'ppid', 'name', 'cmdline']):
            try:
                if proc.info['name'] == 'pt_main_thread':
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    
                    # 跳过主服务进程
                    if proc.info['pid'] == current_pid:
                        print(f"安全清理：跳过主服务进程 {proc.info['pid']}")
                        continue
                    
                    # 跳过包含app.py或start_app.py的进程（主服务进程）
                    if 'app.py' in cmdline or 'start_app.py' in cmdline:
                        print(f"安全清理：跳过主服务进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                        continue
                    
                    # 杀死所有其他pt_main_thread进程
                    print(f"安全清理：发现pt_main_thread进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                    proc.terminate()
                    killed_count += 1
                    try:
                        proc.wait(timeout=1)
                    except psutil.TimeoutExpired:
                        print(f"pt_main_thread进程 {proc.info['pid']} 未响应terminate，强制杀死")
                        proc.kill()
                        try:
                            proc.wait(timeout=1)
                        except psutil.TimeoutExpired:
                            print(f"pt_main_thread进程 {proc.info['pid']} 无法被杀死")
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        if killed_count > 0:
            print(f"安全清理：已终止 {killed_count} 个pt_main_thread进程")
        
        return killed_count
    except Exception as e:
        print(f"安全清理pt_main_thread进程时出错: {e}")
        return 0

def get_system_resources():
    """获取系统资源使用情况"""
    with _SYSTEM_RESOURCES_LOCK:
        return SYSTEM_RESOURCES.copy()

def update_system_resources():
    """更新系统资源统计"""
    try:
        import psutil
        
        # 更新CPU和内存使用率
        cpu_percent = psutil.cpu_percent(interval=0.1)
        memory = psutil.virtual_memory()
        
        # 更新GPU内存使用率
        gpu_memory = 0.0
        try:
            import torch
            if torch.cuda.is_available():
                gpu_memory = torch.cuda.memory_allocated() / (1024**3)  # GB
        except:
            pass
        
        with _SYSTEM_RESOURCES_LOCK:
            SYSTEM_RESOURCES['cpu_usage'] = cpu_percent
            SYSTEM_RESOURCES['memory_usage'] = memory.percent
            SYSTEM_RESOURCES['gpu_memory_usage'] = gpu_memory
            SYSTEM_RESOURCES['active_tasks'] = len(TASK_TIMEOUTS)
            SYSTEM_RESOURCES['active_users'] = len(USER_SESSIONS)
        
        return SYSTEM_RESOURCES.copy()
    except Exception as e:
        print(f"更新系统资源时出错: {e}")
        return SYSTEM_RESOURCES.copy()

def cleanup_inactive_users():
    """清理非活跃用户会话"""
    current_time = time.time()
    inactive_threshold = 3600  # 1小时无活动则清理
    cleaned_count = 0
    
    with _USER_SESSIONS_LOCK:
        inactive_users = []
        for user_id, session in USER_SESSIONS.items():
            if current_time - session['last_activity'] > inactive_threshold:
                inactive_users.append(user_id)
        
        for user_id in inactive_users:
            # 取消该用户的所有任务
            user_tasks = list(session['tasks'])
            for task_id in user_tasks:
                mark_task_cancelled(task_id)
                with _TASK_TIMEOUTS_LOCK:
                    TASK_TIMEOUTS.pop(task_id, None)
                with _TASK_TO_USER_LOCK:
                    TASK_TO_USER.pop(task_id, None)
                terminate_task_processes(task_id)
                clear_task_processes(task_id)
            
            # 删除用户会话
            del USER_SESSIONS[user_id]
            cleaned_count += 1
            print(f"清理非活跃用户: {user_id} (IP: {session['ip']})")
    
    if cleaned_count > 0:
        print(f"已清理 {cleaned_count} 个非活跃用户会话")
        # 更新系统资源统计
        with _SYSTEM_RESOURCES_LOCK:
            SYSTEM_RESOURCES['active_users'] = len(USER_SESSIONS)
    
    return cleaned_count

def start_user_cleanup_checker():
    """启动用户清理检查器"""
    global _user_cleanup_thread, _user_cleanup_running
    
    with _user_cleanup_lock:
        if _user_cleanup_running:
            return
        
        _user_cleanup_running = True
        _user_cleanup_thread = threading.Thread(target=_user_cleanup_worker, daemon=True)
        _user_cleanup_thread.start()
        print("用户清理检查器已启动")

def stop_user_cleanup_checker():
    """停止用户清理检查器"""
    global _user_cleanup_running
    
    with _user_cleanup_lock:
        _user_cleanup_running = False
        print("用户清理检查器已停止")

def _user_cleanup_worker():
    """用户清理工作线程"""
    global _user_cleanup_running
    
    while _user_cleanup_running:
        try:
            # 每5分钟清理一次非活跃用户
            time.sleep(300)
            if _user_cleanup_running:
                cleaned_count = cleanup_inactive_users()
                if cleaned_count > 0:
                    print(f"定期清理：已清理 {cleaned_count} 个非活跃用户")
                # 检查长时间隐藏页面的用户任务
                cleanup_hidden_page_tasks()
        except Exception as e:
            print(f"用户清理检查器出错: {e}")
            time.sleep(60)  # 出错后等待1分钟再继续

def cleanup_hidden_page_tasks():
    """清理长时间隐藏页面的用户任务"""
    try:
        current_time = time.time()
        hidden_threshold = 60  # 页面隐藏超过60秒的任务将被取消
        
        with _USER_SESSIONS_LOCK:
            for user_id, session_info in list(USER_SESSIONS.items()):
                # 检查用户是否有隐藏页面标记
                if session_info.get('page_hidden', False):
                    hidden_time = session_info.get('page_hidden_time', 0)
                    if current_time - hidden_time > hidden_threshold:
                        # 取消该用户的所有任务
                        user_tasks = list(session_info['tasks'])
                        if user_tasks:
                            print(f"用户 {user_id} 页面隐藏超过{hidden_threshold}秒，取消 {len(user_tasks)} 个任务")
                            for task_id in user_tasks:
                                mark_task_cancelled(task_id)
                            # 清理用户会话
                            del USER_SESSIONS[user_id]
    except Exception as e:
        print(f"清理隐藏页面任务时出错: {e}")

def mark_user_page_hidden(user_id):
    """标记用户页面为隐藏状态"""
    with _USER_SESSIONS_LOCK:
        if user_id in USER_SESSIONS:
            USER_SESSIONS[user_id]['page_hidden'] = True
            USER_SESSIONS[user_id]['page_hidden_time'] = time.time()

def mark_user_page_visible(user_id):
    """标记用户页面为可见状态"""
    with _USER_SESSIONS_LOCK:
        if user_id in USER_SESSIONS:
            USER_SESSIONS[user_id]['page_hidden'] = False
            USER_SESSIONS[user_id].pop('page_hidden_time', None)

def cleanup_old_cancelled_tasks():
    """清理过期的取消任务（防止内存泄漏）"""
    # 这里可以添加时间戳机制来清理过期的任务
    # 暂时保持简单，在任务完成时清理
    pass

def start_task_timer(task_id):
    """开始任务计时"""
    if task_id:
        with _TASK_TIMEOUTS_LOCK:
            TASK_TIMEOUTS[task_id] = time.time()
            print(f"任务 {task_id} 开始计时")
            # 如果这是第一个任务，启动超时检查器
            if len(TASK_TIMEOUTS) == 1:
                start_timeout_checker()

def check_task_timeout(task_id):
    """检查任务是否超时"""
    if not task_id:
        return False
    
    with _TASK_TIMEOUTS_LOCK:
        if task_id in TASK_TIMEOUTS:
            start_time = TASK_TIMEOUTS[task_id]
            elapsed = time.time() - start_time
            if elapsed > TASK_TIMEOUT_SECONDS:
                print(f"任务 {task_id} 超时 ({elapsed:.1f}秒)，自动取消并终止进程")
                mark_task_cancelled(task_id)
                # 强制终止所有相关进程
                terminate_task_processes(task_id)
                return True
    return False

def clear_task_timer(task_id):
    """清除任务计时"""
    if task_id:
        with _TASK_TIMEOUTS_LOCK:
            TASK_TIMEOUTS.pop(task_id, None)
            print(f"任务 {task_id} 计时已清除")
            # 如果没有任务了，停止超时检查器
            if len(TASK_TIMEOUTS) == 0:
                stop_timeout_checker()

def register_active_process(task_id, process):
    """注册活跃进程"""
    if task_id and process:
        with _ACTIVE_PROCESSES_LOCK:
            if task_id not in ACTIVE_PROCESSES:
                ACTIVE_PROCESSES[task_id] = []
            ACTIVE_PROCESSES[task_id].append(process)
            print(f"任务 {task_id} 注册了进程 {process.pid}")

def terminate_task_processes(task_id):
    """强制终止任务的所有进程"""
    if not task_id:
        return
    
    with _ACTIVE_PROCESSES_LOCK:
        if task_id in ACTIVE_PROCESSES:
            processes = ACTIVE_PROCESSES[task_id]
            for process in processes:
                try:
                    if process.is_alive():
                        print(f"强制终止进程 {process.pid}")
                        process.terminate()
                        process.join(timeout=5)  # 等待5秒
                        if process.is_alive():
                            print(f"进程 {process.pid} 未响应，强制杀死")
                            process.kill()
                except Exception as e:
                    print(f"终止进程时出错: {e}")
            del ACTIVE_PROCESSES[task_id]
            print(f"任务 {task_id} 的所有进程已终止")

def clear_task_processes(task_id):
    """清除任务进程记录"""
    if task_id:
        with _ACTIVE_PROCESSES_LOCK:
            ACTIVE_PROCESSES.pop(task_id, None)

def _timeout_checker_worker():
    """超时检查工作线程"""
    global _timeout_checker_running
    while _timeout_checker_running:
        try:
            with _TASK_TIMEOUTS_LOCK:
                current_time = time.time()
                expired_tasks = []
                for task_id, start_time in TASK_TIMEOUTS.items():
                    if current_time - start_time > TASK_TIMEOUT_SECONDS:
                        expired_tasks.append(task_id)
                
                for task_id in expired_tasks:
                    elapsed_time = current_time - TASK_TIMEOUTS[task_id]
                    print(f"超时检查器发现任务 {task_id} 超时，强制终止 (运行时间: {elapsed_time:.1f}秒, 超时阈值: {TASK_TIMEOUT_SECONDS}秒)")
                    mark_task_cancelled(task_id)
                    TASK_TIMEOUTS.pop(task_id, None)
                    # 强制清理该任务的所有进程
                    terminate_task_processes(task_id)
                    # 额外清理所有multiprocessing进程
                    kill_all_multiprocessing_processes()
                    # 强制清理所有Python进程
                    force_kill_all_python_processes()
            
            time.sleep(5)  # 每5秒检查一次
        except Exception as e:
            print(f"超时检查器出错: {e}")
            time.sleep(5)

def start_timeout_checker():
    """启动超时检查线程"""
    global _timeout_checker_thread, _timeout_checker_running
    
    with _timeout_checker_lock:
        if not _timeout_checker_running:
            _timeout_checker_running = True
            _timeout_checker_thread = threading.Thread(target=_timeout_checker_worker, daemon=True)
            _timeout_checker_thread.start()
            print("超时检查器已启动")

def stop_timeout_checker():
    """停止超时检查线程"""
    global _timeout_checker_running
    
    with _timeout_checker_lock:
        _timeout_checker_running = False
        print("超时检查器已停止")

def force_cleanup_all_processes():
    """强制清理所有活跃进程"""
    with _ACTIVE_PROCESSES_LOCK:
        for task_id, processes in list(ACTIVE_PROCESSES.items()):
            print(f"强制清理任务 {task_id} 的 {len(processes)} 个进程")
            for process in processes:
                try:
                    if process.is_alive():
                        print(f"强制终止进程 {process.pid}")
                        process.terminate()
                        process.join(timeout=2)
                        if process.is_alive():
                            print(f"进程 {process.pid} 未响应，强制杀死")
                            process.kill()
                except Exception as e:
                    print(f"清理进程时出错: {e}")
        ACTIVE_PROCESSES.clear()
        print("所有进程已强制清理")

def kill_all_multiprocessing_processes():
    """智能杀死应用相关的multiprocessing进程，保护系统重要进程"""
    try:
        import psutil
        current_pid = os.getpid()
        killed_count = 0
        
        for proc in psutil.process_iter(['pid', 'name', 'cmdline', 'ppid']):
            try:
                if proc.info['name'] == 'python' and proc.info['pid'] != current_pid:
                    cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                    
                    # 更精确的进程识别：只杀死我们应用相关的进程
                    should_kill = False
                    
                    # 1. 检查是否是我们应用的工作进程
                    if 'spawn_main' in cmdline and any(keyword in cmdline for keyword in [
                        'storage_compare', 'keyword_search', 'single_compare', 'doc_similarity'
                    ]):
                        should_kill = True
                        print(f"发现应用工作进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                    
                    # 2. 跳过重要的系统进程（resource_tracker等）
                    elif 'resource_tracker' in cmdline:
                        print(f"跳过系统资源跟踪进程 {proc.info['pid']}")
                        should_kill = False
                    
                    # 3. 检查是否是我们应用的直接子进程
                    elif proc.info['ppid'] == current_pid and 'multiprocessing' in cmdline:
                        should_kill = True
                        print(f"发现应用子进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                    
                    if should_kill:
                        proc.terminate()
                        killed_count += 1
                        # 等待进程终止
                        try:
                            proc.wait(timeout=2)
                        except psutil.TimeoutExpired:
                            print(f"进程 {proc.info['pid']} 未响应terminate，强制杀死")
                            proc.kill()
                            
            except (psutil.NoSuchProcess, psutil.AccessDenied, psutil.TimeoutExpired):
                pass
        
        print(f"已杀死 {killed_count} 个multiprocessing进程")
        return killed_count
    except Exception as e:
        print(f"杀死multiprocessing进程时出错: {e}")
        return 0

def force_kill_all_python_processes():
    """智能杀死应用相关的Python进程，保护系统重要进程"""
    try:
        import psutil
        current_pid = os.getpid()
        killed_count = 0
        
        # 进行多次清理，确保所有进程都被终止
        for attempt in range(3):
            print(f"第 {attempt + 1} 次清理Python进程...")
            current_killed = 0
            
            for proc in psutil.process_iter(['pid', 'name', 'cmdline', 'ppid']):
                try:
                    if proc.info['name'] == 'python' and proc.info['pid'] != current_pid:
                        cmdline = ' '.join(proc.info['cmdline']) if proc.info['cmdline'] else ''
                        
                        # 更精确的进程识别：只杀死我们应用相关的进程
                        should_kill = False
                        
                        # 1. 检查是否是我们应用的工作进程
                        if any(keyword in cmdline for keyword in ['storage_compare', 'keyword_search', 'single_compare', 'doc_similarity']):
                            should_kill = True
                        
                        # 2. 检查是否是spawn_main进程（但跳过resource_tracker）
                        elif 'spawn_main' in cmdline and 'resource_tracker' not in cmdline:
                            should_kill = True
                        
                        # 3. 跳过重要的系统进程
                        elif any(system_keyword in cmdline for system_keyword in ['resource_tracker', 'forkserver', 'semlock_tracker']):
                            print(f"跳过系统进程 {proc.info['pid']}: {cmdline[:100]}...")
                            should_kill = False
                        
                        # 4. 检查是否是我们应用的直接子进程
                        elif proc.info['ppid'] == current_pid and 'multiprocessing' in cmdline:
                            should_kill = True
                        
                        if should_kill:
                            print(f"强制杀死Python进程 {proc.info['pid']}，命令行: {cmdline[:100]}...")
                            try:
                                proc.kill()  # 直接使用kill，不等待
                                current_killed += 1
                            except psutil.NoSuchProcess:
                                pass  # 进程已经不存在
                                
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    pass
            
            killed_count += current_killed
            print(f"第 {attempt + 1} 次清理杀死了 {current_killed} 个进程")
            
            if current_killed == 0:
                break  # 没有进程需要清理了
            
            time.sleep(1)  # 等待1秒再进行下一次清理
        
        print(f"总共杀死了 {killed_count} 个Python进程")
        return killed_count
    except Exception as e:
        print(f"强制杀死Python进程时出错: {e}")
        return 0

# 多核处理配置
CPU_CORES = multiprocessing.cpu_count()

# 只在主进程中打印初始化信息
if is_main_initialization_process():
    print(f"检测到CPU核心数: {CPU_CORES}")
MAX_WORKERS = min(CPU_CORES, 8)  # 限制最大工作线程数
if is_main_initialization_process():
    print(f"使用工作线程数: {MAX_WORKERS}")

# 只在主进程中打印初始化信息
if is_main_initialization_process():
    print("\n" + "⏳" + "="*58 + "⏳")
    print("🔧 系统正在初始化...")
    print("📋 正在检测硬件配置和加载AI模型")
    print("⏱️  这可能需要10-30秒，请耐心等待")
    print("💡 首次运行或网络较慢时可能需要更长时间")
    print("⏳" + "="*58 + "⏳")

# 确保目录存在
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(os.path.join(UPLOAD_FOLDER, 'temp'), exist_ok=True)
os.makedirs(STORAGE_FOLDER, exist_ok=True)

# 加载模型，支持GPU与本地离线
def get_device():
    """获取最佳可用设备"""
    # 使用更严格的主进程检测
    should_log = is_main_initialization_process()
    
    if should_log:
        print("🔍 开始硬件设备检测...")
        print("  ├─ 检测CUDA支持...")
    
    if torch.cuda.is_available():
        if should_log:
            print("  ├─ ✅ CUDA可用")
            cuda_version = torch.version.cuda
            print(f"  ├─ 📋 CUDA版本: {cuda_version}")
            print("  ├─ 🔧 配置GPU设备...")
        
        torch.cuda.set_device(0)
        device = 'cuda:0'
        
        if should_log:
            gpu_name = torch.cuda.get_device_name(0)
            gpu_memory = torch.cuda.get_device_properties(0).total_memory / 1024**3
            print(f"  ├─ 🎮 GPU设备: {gpu_name}")
            print(f"  ├─ 💾 GPU内存: {gpu_memory:.1f} GB")
            print("  └─ ✅ GPU设备配置完成")
    else:
        if should_log:
            print("  ├─ ❌ CUDA不可用")
            print("  ├─ 🔄 切换到CPU模式")
        device = 'cpu'
        if should_log:
            print("  └─ ✅ CPU设备配置完成")
    
    return device

print("🔧 [DEBUG] 开始获取设备信息...")
device = get_device()
print(f"🔧 [DEBUG] 设备获取完成: {device}")

# 醒目显示硬件检测信息（只在主进程中）
print("🔧 [DEBUG] 准备显示硬件检测信息...")
if is_main_initialization_process():
    print("\n" + "="*60)
    if device.startswith('cuda'):
        print("🔍 GPU硬件检测")
        print(f"🖥️  GPU设备: {torch.cuda.get_device_name(0)}")
        print(f"💾 GPU内存: {torch.cuda.get_device_properties(0).total_memory / 1024**3:.1f} GB")
        print("✅ GPU可用，优先为模型使用GPU加速")
    else:
        print("🔍 CPU硬件检测")
        print(f"⚙️  CPU核心数: {CPU_CORES}")
        print(f"🔧 工作线程数: {MAX_WORKERS}")
        print("💡 将使用CPU多线程计算")
    print(f"🎯 默认设备: {device}")
    print("="*60)

# 准备后备：当无法在线下载模型时，使用TF-IDF
_model = None
_use_tfidf_fallback = False
_model_source = None  # 记录模型来源：'online_download', 'local_cache', 'local_path', 'tfidf_fallback', 'fast_start'
_startup_mode = None  # 记录启动模式：'normal', 'fast_start', 'force_tfidf'
_model_load_time = 0  # 记录模型加载时间

# 允许通过环境变量指定本地模型目录或镜像名称
MODEL_NAME_OR_PATH = os.environ.get('ST_MODEL_PATH', 'all-mpnet-base-v2')
LOCAL_FILES_ONLY = os.environ.get('HF_LOCAL_ONLY', '0') == '1'
HF_MIRROR = os.environ.get('HF_ENDPOINT')  # 如 https://hf-mirror.com

# --- 智能启动逻辑（在模块导入时执行）---

print("🔧 [DEBUG] 检查是否为主初始化进程...")
# 只在主进程中执行启动逻辑，避免在子进程中重复初始化
if is_main_initialization_process():
    print("🔧 [DEBUG] 确认为主进程，开始执行启动逻辑")
    print("\n🚀 开始智能启动流程...")
    
    try:
        print("🔧 [DEBUG] 准备导入启动管理器...")
        # 导入启动管理器
        from startup_manager import run_intelligent_startup
        print("🔧 [DEBUG] 启动管理器导入成功")
        
        print("🔧 [DEBUG] 开始运行智能启动流程...")
        # 运行智能启动流程
        startup_result = run_intelligent_startup()
        print("🔧 [DEBUG] 智能启动流程完成")
        
        # 3. 更新全局变量
        if startup_result['success']:
            execution_result = startup_result['execution_result']
            _model = execution_result['model']
            actual_config = execution_result['actual_config']
            
            # 设置模型来源和启动模式
            strategy_to_source = {
                'local_model': 'local_cache',
                'download_model': 'online_download', 
                'tfidf': 'tfidf_fallback',
                'minimal': 'basic_fallback'
            }
            
            _model_source = strategy_to_source.get(actual_config['strategy'], 'unknown')
            _startup_mode = 'normal' if _model is not None else 'fast_start'
            _model_load_time = actual_config.get('load_time', 0)
            
            # 配置TF-IDF
            if actual_config['strategy'] in ['tfidf', 'minimal']:
                _use_tfidf_fallback = True
                try:
                    from sklearn.feature_extraction.text import TfidfVectorizer
                    from sklearn.metrics.pairwise import cosine_similarity as _sk_cos_sim
                    _tfidf = TfidfVectorizer(max_features=50000, ngram_range=(1, 2))
                    print("✅ TF-IDF备用方案已初始化")
                except Exception as e:
                    print(f"⚠️ TF-IDF初始化警告: {e}")
                    _tfidf = None
            else:
                _use_tfidf_fallback = False
                _tfidf = None
            
            print(f"\n✅ 启动完成!")
            print(f"📋 最终配置: {actual_config}")
            
        else:
            raise Exception(f"启动策略执行失败: {startup_result.get('error', '未知错误')}")
            
    except Exception as e:
        print(f"❌ 智能启动失败: {e}")
        print("🔄 执行紧急后备启动...")
        
        # 紧急后备：最简单的TF-IDF模式
        _model = None
        _model_source = 'emergency_fallback'
        _startup_mode = 'emergency'
        _model_load_time = 0
        _use_tfidf_fallback = True
        
        try:
            from sklearn.feature_extraction.text import TfidfVectorizer
            from sklearn.metrics.pairwise import cosine_similarity as _sk_cos_sim
            _tfidf = TfidfVectorizer(max_features=10000, ngram_range=(1, 1))
            print("✅ 紧急TF-IDF模式启动成功")
        except Exception:
            print("❌ 紧急启动也失败，某些功能可能不可用")
            _tfidf = None
else:
    # 在子进程中，设置默认值
    _model = None
    _use_tfidf_fallback = True
    _model_source = 'subprocess_default'
    _startup_mode = 'subprocess'
    _model_load_time = 0
    _tfidf = None

def find_cached_model(model_name):
    """查找缓存中的模型目录"""
    # 使用更严格的主进程检测
    should_log = is_main_initialization_process()
    
    def log_if_main(message):
        """只在主进程中打印日志"""
        if should_log:
            print(message)
    
    log_if_main("🔍 搜索本地缓存模型...")
    
    possible_cache_dirs = [
        os.path.expanduser('~/.cache/huggingface/hub'),
        os.path.expanduser('~/.cache/huggingface/transformers'),
        '/data/doc_similarity_env/cache/huggingface/hub',
        '/root/.cache/huggingface/hub',
        os.path.join(os.getcwd(), '.cache', 'huggingface', 'hub')  # 当前目录下的缓存
    ]
    
    for i, cache_dir in enumerate(possible_cache_dirs, 1):
        log_if_main(f"  ├─ 检查目录 {i}/{len(possible_cache_dirs)}: {cache_dir}")
        
        if os.path.exists(cache_dir):
            log_if_main("  │  ├─ ✅ 目录存在")
            try:
                items = os.listdir(cache_dir)
                log_if_main(f"  │  ├─ 📁 发现 {len(items)} 个项目")
                
                # 查找包含模型名称的目录
                found_models = []
                for item in items:
                    if model_name.replace('/', '--') in item or model_name in item:
                        found_models.append(item)
                
                if found_models:
                    print(f"  │  ├─ 🎯 找到 {len(found_models)} 个相关模型:")
                    for model in found_models:
                        print(f"  │  │  └─ {model}")
                    
                    for model in found_models:
                        model_path = os.path.join(cache_dir, model)
                        print(f"  │  ├─ 🔍 验证模型: {model}")
                        
                        if os.path.isdir(model_path):
                            print("  │  │  ├─ ✅ 是目录")
                            # 查找包含config.json或pytorch_model.bin的子目录
                            for root, dirs, files in os.walk(model_path):
                                if any(f in files for f in ['config.json', 'pytorch_model.bin', 'model.safetensors']):
                                    print(f"  │  │  ├─ ✅ 找到有效模型文件")
                                    print(f"  │  │  └─ 📍 路径: {root}")
                                    print("  └─ 🎉 模型验证成功!")
                                    return root
                            print("  │  │  └─ ❌ 未找到有效模型文件")
                        else:
                            print("  │  │  └─ ❌ 不是目录")
                else:
                    print("  │  └─ ❌ 未找到相关模型")
                    
            except PermissionError as e:
                print(f"  │  └─ ❌ 权限不足: {e}")
                continue
            except Exception as e:
                print(f"  │  └─ ❌ 检查失败: {e}")
                continue
        else:
            print("  │  └─ ❌ 目录不存在")
    
    print("  └─ 💔 未找到任何缓存模型")
    return None

def create_simple_embedding_model():
    """创建一个简单的基于TF-IDF的嵌入模型作为后备"""
    try:
        from sklearn.feature_extraction.text import TfidfVectorizer
        from sklearn.metrics.pairwise import cosine_similarity
        import numpy as np
        
        class SimpleTfidfModel:
            def __init__(self):
                self.vectorizer = TfidfVectorizer(
                    max_features=10000,
                    ngram_range=(1, 2),
                    stop_words='english' if 'english' in str(self) else None
                )
                self.device = 'cpu'
                # 只在主进程中打印TF-IDF模型创建信息
                if is_main_initialization_process():
                    print("创建了简单的TF-IDF嵌入模型")
            
            def encode(self, sentences, convert_to_tensor=False, **kwargs):
                if isinstance(sentences, str):
                    sentences = [sentences]
                
                # 拟合并转换文本
                vectors = self.vectorizer.fit_transform(sentences)
                
                if convert_to_tensor:
                    # 转换为torch tensor
                    import torch
                    return torch.tensor(vectors.toarray(), dtype=torch.float32)
                else:
                    return vectors.toarray()
        
        return SimpleTfidfModel()
    except ImportError:
        print("无法创建TF-IDF后备模型，缺少sklearn")
        return None
    except Exception as e:
        print(f"创建TF-IDF后备模型失败: {e}")
        return None

def load_sentence_transformer_with_timeout(model_path, load_kwargs, timeout=30):
    """带超时的模型加载"""
    import signal
    import threading
    
    result = {'model': None, 'error': None}
    
    def load_model():
        try:
            from sentence_transformers import SentenceTransformer
            # 在子进程中抑制模型加载日志
            if not is_main_initialization_process():
                import warnings
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    result['model'] = SentenceTransformer(model_path, **load_kwargs)
            else:
                result['model'] = SentenceTransformer(model_path, **load_kwargs)
        except Exception as e:
            result['error'] = e
    
    # 创建加载线程
    thread = threading.Thread(target=load_model)
    thread.daemon = True
    thread.start()
    thread.join(timeout=timeout)
    
    if thread.is_alive():
        print(f"模型加载超时 ({timeout}秒)，可能是网络问题或模型文件损坏")
        return None, "加载超时"
    
    return result['model'], result['error']

def check_internet_connection(url="https://huggingface.co", timeout=2):
    """检查是否能访问Hugging Face"""
    should_log = is_main_initialization_process()
    
    if should_log:
        print("🌐 检测网络连接...")
        print(f"  ├─ 目标地址: {url}")
        print(f"  ├─ 超时设置: {timeout}秒")
        print("  ├─ 🔗 尝试连接...")
    
    try:
        import urllib.request
        urllib.request.urlopen(url, timeout=timeout)
        if should_log:
            print("  └─ ✅ 网络连接成功")
        return True
    except Exception as e:
        if should_log:
            print(f"  └─ ❌ 网络连接失败: {type(e).__name__}")
        return False

def print_download_instructions():
    """打印下载说明"""
    print("\n" + "="*80)
    print("📋 模型下载说明")
    print("="*80)
    print("由于网络无法访问Hugging Face，请手动下载模型文件：")
    print()
    print("🔗 下载地址:")
    print("  方法1 - Hugging Face官方:")
    print("    https://huggingface.co/sentence-transformers/all-mpnet-base-v2")
    print("  方法2 - 镜像站点:")
    print("    https://hf-mirror.com/sentence-transformers/all-mpnet-base-v2")
    print("  方法3 - ModelScope:")
    print("    https://modelscope.cn/models/sentence-transformers/all-mpnet-base-v2")
    print()
    print("📁 存放路径:")
    cache_dir = os.path.expanduser('~/.cache/huggingface/hub')
    model_dir = os.path.join(cache_dir, 'models--sentence-transformers--all-mpnet-base-v2', 'snapshots', 'latest')
    print(f"    {model_dir}")
    print()
    print("📦 需要的文件:")
    print("  - config.json")
    print("  - pytorch_model.bin 或 model.safetensors")
    print("  - tokenizer.json")
    print("  - tokenizer_config.json")
    print("  - vocab.txt")
    print("  - sentence_bert_config.json")
    print("  - config_sentence_transformers.json")
    print()
    print("💡 下载完成后，重新启动应用即可自动加载模型")
    print("="*80)

def auto_download_model(model_name="sentence-transformers/all-mpnet-base-v2"):
    """自动下载模型"""
    should_log = is_main_initialization_process()
    
    if should_log:
        print(f"📥 开始自动下载模型...")
        print(f"  ├─ 模型名称: {model_name}")
        print("  ├─ 🔧 配置下载环境...")
    
    try:
        from sentence_transformers import SentenceTransformer
        # 临时关闭离线模式以允许下载
        os.environ.pop('TRANSFORMERS_OFFLINE', None)
        os.environ.pop('HF_HUB_OFFLINE', None)
        if should_log:
            print("  ├─ ✅ 环境配置完成")
        
        if should_log:
            print("  ├─ ⏳ 开始下载模型文件...")
            print("  │  ├─ 这可能需要几分钟时间")
            print("  │  └─ 请保持网络连接稳定")
        
        model = SentenceTransformer(model_name, device=device)
        
        if should_log:
            print("  ├─ ✅ 模型下载完成")
            print("  └─ 🎉 模型加载成功!")
        return model
    except Exception as e:
        print(f"  └─ ❌ 自动下载失败: {e}")
        return None

def load_model_with_timeout(model_name, timeout=45):
    """
    在独立的进程中加载模型，并设置超时保护。
    这是解决模型加载卡死的关键。
    """
    import time
    from multiprocessing import Process, Queue

    def worker(q, model_name_to_load, device):
        """在子进程中运行的加载函数"""
        try:
            # 在子进程中，日志是不可见的，但加载仍在进行
            from sentence_transformers import SentenceTransformer
            model = SentenceTransformer(model_name_to_load, device=device)
            q.put(model)
        except Exception as e:
            # 将异常信息放入队列，以便主进程可以捕获
            q.put(e)

    q = Queue()
    p = Process(target=worker, args=(q, model_name, device))
    p.daemon = True
    
    print(f"🔄 开始加载SentenceTransformer模型: '{model_name}' (超时: {timeout}秒)...")
    start_time = time.time()
    
    p.start()
    p.join(timeout)
    
    elapsed_time = time.time() - start_time

    if p.is_alive():
        print(f"❌ 模型加载超时！(超过 {timeout} 秒)")
        p.terminate()  # 强制终止子进程
        p.join()
        return None, "timeout", elapsed_time

    if not q.empty():
        result = q.get()
        if isinstance(result, Exception):
            print(f"❌ 模型加载时发生错误: {result}")
            return None, str(result), elapsed_time
        else:
            # 成功加载
            # 检查模型是否真的在GPU上 (如果适用)
            model_device = "CPU"
            if torch.cuda.is_available():
                try:
                    # 尝试将一个测试张量放到模型设备上，验证可用性
                    test_tensor = torch.tensor([1]).to(result.device)
                    model_device = f"GPU ({result.device})"
                except Exception:
                    model_device = "CPU (GPU验证失败)"
            
            print(f"✅ 模型加载成功! (耗时: {elapsed_time:.2f}秒, 运行于: {model_device})")
            return result, "success", elapsed_time
    
    return None, "unknown_error", elapsed_time

def allowed_file(filename):
    """检查文件扩展名"""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def clean_text(text):
    """清理文本：去除噪声、页眉页脚，保留格式"""
    if not text:
        return text
    text = re.sub(r'Page \d+ of \d+|Confidential|© \d{4}.*?\|', '', text)
    return text.strip()

def extract_formatted_text(file_path, is_storage_file=False, task_id=None):
    """提取文件文本并保留格式，返回行列表、HTML格式内容和临时文件路径"""
    try:
        ext = file_path.rsplit('.', 1)[1].lower()
        lines = []
        html_content = []
        temp_file_path = file_path  # 默认不移动文件

        # 如果是上传文件，复制到临时目录；如果是存储库文件，直接读取
        if not is_storage_file:
            temp_file_path = os.path.join(UPLOAD_FOLDER, 'temp', f"{uuid.uuid4()}.{ext}")
            os.makedirs(os.path.dirname(temp_file_path), exist_ok=True)
            shutil.copy(file_path, temp_file_path)  # 复制而非移动

        if ext in {'yaml', 'yml'}:
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                data = yaml.safe_load(f)
                text = yaml.dump(data, allow_unicode=True, sort_keys=False, default_flow_style=False)
                lines = text.splitlines()
                html_content = [f"<pre>{line}</pre>" for line in lines]

        elif ext in {'xlsx', 'xls'}:
            try:
                print(f"开始处理Excel文件: {temp_file_path}")
                
                # 检查任务是否被取消
                if is_task_cancelled(task_id):
                    print(f"任务 {task_id} 已被取消，停止处理Excel文件")
                    return None, None, None
                
                if ext == 'xlsx':
                    # 对于xlsx文件，使用openpyxl引擎
                    print("使用openpyxl引擎处理xlsx文件")
                    df = pd.read_excel(temp_file_path, engine='openpyxl')
                else:
                    # 对于xls文件，尝试多种方法
                    df = None
                    
                    # 方法1：尝试xlrd引擎
                    try:
                        print("尝试使用xlrd引擎处理xls文件")
                        df = pd.read_excel(temp_file_path, engine='xlrd')
                        print("xlrd引擎成功")
                    except Exception as e:
                        print(f"xlrd引擎失败: {e}")
                        
                        # 检查任务是否被取消
                        if is_task_cancelled(task_id):
                            print(f"任务 {task_id} 已被取消，停止处理Excel文件")
                            return None, None, None
                        
                        # 方法2：尝试openpyxl引擎
                        try:
                            print("尝试使用openpyxl引擎处理xls文件")
                            df = pd.read_excel(temp_file_path, engine='openpyxl')
                            print("openpyxl引擎成功")
                        except Exception as e:
                            print(f"openpyxl引擎失败: {e}")
                            
                            # 检查任务是否被取消
                            if is_task_cancelled(task_id):
                                print(f"任务 {task_id} 已被取消，停止处理Excel文件")
                                return None, None, None
                            
                            # 方法3：尝试不指定引擎
                            try:
                                print("尝试使用默认引擎处理xls文件")
                                df = pd.read_excel(temp_file_path)
                                print("默认引擎成功")
                            except Exception as e:
                                print(f"默认引擎失败: {e}")
                                
                                # 方法4：尝试使用xlwings（仅Windows环境）
                                if os.name == 'nt':  # 只在Windows环境下尝试xlwings
                                    try:
                                        import xlwings as xw
                                        print("尝试使用xlwings处理xls文件")
                                        app = xw.App(visible=False)
                                        wb = app.books.open(temp_file_path)
                                        ws = wb.sheets[0]
                                        data = ws.used_range.value
                                        wb.close()
                                        app.quit()
                                        
                                        if data:
                                            # 将xlwings数据转换为DataFrame
                                            df = pd.DataFrame(data[1:], columns=data[0])
                                            print("xlwings成功")
                                    except ImportError:
                                        print("xlwings不可用")
                                    except Exception as e:
                                        print(f"xlwings失败: {e}")
                                else:
                                    print("Linux环境跳过xlwings")
                                
                                # 方法5：尝试使用pyexcel（跨平台）
                                if df is None:
                                    try:
                                        import pyexcel as pe
                                        print("尝试使用pyexcel处理xls文件")
                                        sheet = pe.get_sheet(file_name=temp_file_path)
                                        data = sheet.to_array()
                                        if data:
                                            # 将pyexcel数据转换为DataFrame
                                            df = pd.DataFrame(data[1:], columns=data[0])
                                            print("pyexcel成功")
                                    except ImportError:
                                        print("pyexcel不可用")
                                    except Exception as e:
                                        print(f"pyexcel失败: {e}")
                                
                                # 方法6：尝试使用xlutils（需要xlrd配合）
                                if df is None:
                                    try:
                                        import xlutils
                                        from xlutils.copy import copy
                                        from xlrd import open_workbook
                                        print("尝试使用xlutils处理xls文件")
                                        rb = open_workbook(temp_file_path, formatting_info=True)
                                        # 读取第一个工作表
                                        sheet = rb.sheet_by_index(0)
                                        data = []
                                        for row_idx in range(sheet.nrows):
                                            row_data = []
                                            for col_idx in range(sheet.ncols):
                                                cell_value = sheet.cell_value(row_idx, col_idx)
                                                row_data.append(cell_value)
                                            data.append(row_data)
                                        
                                        if data:
                                            df = pd.DataFrame(data[1:], columns=data[0])
                                            print("xlutils成功")
                                    except ImportError:
                                        print("xlutils不可用")
                                    except Exception as e:
                                        print(f"xlutils失败: {e}")
                                
                                # 方法7：尝试使用LibreOffice命令行工具（Linux环境）
                                if df is None and os.name != 'nt':
                                    try:
                                        import subprocess
                                        import tempfile
                                        print("尝试使用LibreOffice命令行工具处理xls文件")
                                        
                                        # 创建临时目录
                                        with tempfile.TemporaryDirectory() as temp_dir:
                                            # 使用LibreOffice转换为CSV
                                            cmd = [
                                                'libreoffice', '--headless', '--convert-to', 'csv',
                                                '--outdir', temp_dir, temp_file_path
                                            ]
                                            result = subprocess.run(cmd, capture_output=True, text=True, timeout=30)
                                            
                                            if result.returncode == 0:
                                                # 查找生成的CSV文件
                                                csv_files = [f for f in os.listdir(temp_dir) if f.endswith('.csv')]
                                                if csv_files:
                                                    csv_path = os.path.join(temp_dir, csv_files[0])
                                                    df = pd.read_csv(csv_path, encoding='utf-8')
                                                    print("LibreOffice命令行工具成功")
                                    except (ImportError, subprocess.TimeoutExpired, FileNotFoundError) as e:
                                        print(f"LibreOffice命令行工具失败: {e}")
                                    except Exception as e:
                                        print(f"LibreOffice命令行工具失败: {e}")
                
                if df is None or df.empty:
                    print("无法读取Excel文件数据")
                    return None, None, None
                
                print(f"成功读取Excel文件，共{len(df)}行{len(df.columns)}列")
                
                # 再次检查任务是否被取消
                if is_task_cancelled(task_id):
                    print(f"任务 {task_id} 已被取消，停止处理Excel文件")
                    return None, None, None
                
                # 检查文件大小，对大文件进行优化处理
                total_cells = len(df) * len(df.columns)
                is_large_file = total_cells > 100000  # 超过10万个单元格视为大文件
                
                if is_large_file:
                    print(f"检测到大文件（{total_cells}个单元格），使用优化处理策略")
                    
                    # 对于大文件，只处理前1000行和前100列，避免内存和性能问题
                    max_rows = min(1000, len(df))
                    max_cols = min(100, len(df.columns))
                    
                    print(f"大文件优化：处理前{max_rows}行和前{max_cols}列")
                    df_optimized = df.iloc[:max_rows, :max_cols]
                    
                    # 处理数据
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在数据转换前已被取消")
                        return None, None, None
                    
                    text = df_optimized.to_csv(index=False, sep='\t')
                    
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在文本分割前已被取消")
                        return None, None, None
                    
                    lines = text.splitlines()
                    
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在HTML生成前已被取消")
                        return None, None, None
                    
                    html_table = df_optimized.to_html(index=False, border=1, classes='table')
                    html_content = [html_table]
                    
                    print(f"大文件优化处理完成，提取了{len(lines)}行文本（原始文件：{len(df)}行{len(df.columns)}列）")
                else:
                    # 处理数据
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在数据转换前已被取消")
                        return None, None, None
                    
                    text = df.to_csv(index=False, sep='\t')
                    
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在文本分割前已被取消")
                        return None, None, None
                    
                    lines = text.splitlines()
                    
                    if is_task_cancelled(task_id):
                        print(f"任务 {task_id} 在HTML生成前已被取消")
                        return None, None, None
                    
                    html_table = df.to_html(index=False, border=1, classes='table')
                    html_content = [html_table]
                
                print(f"Excel文件处理完成，提取了{len(lines)}行文本")
                
            except Exception as e:
                print(f"Excel文件处理失败: {e}")
                import traceback
                traceback.print_exc()
                return None, None, None

        elif ext == 'docx':
            doc = Document(temp_file_path)
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    lines.append(text)
                    style = para.style.name
                    html_tag = 'p'
                    if style.startswith('Heading'):
                        html_tag = f'h{style[-1]}'
                    elif para.style.name == 'List Bullet' or para.text.strip().startswith(('-', '*')):
                        html_content.append(f'<li>{text}</li>')
                    else:
                        html_content.append(f'<{html_tag}>{text}</{html_tag}>')
            for table in doc.tables:
                html_table = '<table border="1" class="table">'
                for row in table.rows:
                    html_table += '<tr>'
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        lines.append(cell_text)
                        html_table += f'<td>{cell_text}</td>'
                    html_table += '</tr>'
                html_table += '</table>'
                html_content.append(html_table)

        elif ext == 'pdf':
            # 尝试多种PDF提取方法
            extracted = False
            
            # 方法1：使用textract（如果可用）
            try:
                import textract
                text = textract.process(temp_file_path).decode('utf-8')
                cleaned_text = clean_text(text)
                lines = cleaned_text.splitlines()
                html_content = [f"<pre>{line}</pre>" for line in lines]
                extracted = True
                print("使用textract成功提取PDF文本")
            except ImportError:
                print("textract不可用，尝试其他方法")
            except Exception as e:
                print(f"textract提取PDF失败: {e}")
            
            # 方法2：使用pdfminer.six
            if not extracted:
                try:
                    from pdfminer.high_level import extract_text
                    text = extract_text(temp_file_path)
                    cleaned_text = clean_text(text)
                    lines = cleaned_text.splitlines()
                    html_content = [f"<pre>{line}</pre>" for line in lines]
                    extracted = True
                    print("使用pdfminer.six成功提取PDF文本")
                except Exception as e:
                    print(f"pdfminer.six提取PDF失败: {e}")
            
            # 方法3：使用pdfplumber
            if not extracted:
                try:
                    import pdfplumber
                    text = ""
                    with pdfplumber.open(temp_file_path) as pdf:
                        for page in pdf.pages:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                    cleaned_text = clean_text(text)
                    lines = cleaned_text.splitlines()
                    html_content = [f"<pre>{line}</pre>" for line in lines]
                    extracted = True
                    print("使用pdfplumber成功提取PDF文本")
                except ImportError:
                    print("pdfplumber不可用")
                except Exception as e:
                    print(f"pdfplumber提取PDF失败: {e}")
            
            # 方法4：使用PyPDF2
            if not extracted:
                try:
                    import PyPDF2
                    text = ""
                    with open(temp_file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page in pdf_reader.pages:
                            text += page.extract_text() + "\n"
                    cleaned_text = clean_text(text)
                    lines = cleaned_text.splitlines()
                    html_content = [f"<pre>{line}</pre>" for line in lines]
                    extracted = True
                    print("使用PyPDF2成功提取PDF文本")
                except ImportError:
                    print("PyPDF2不可用")
                except Exception as e:
                    print(f"PyPDF2提取PDF失败: {e}")
            
            if not extracted:
                print("所有PDF提取方法都失败")
                return None, None, None
        
        elif ext == 'doc':
            # 尝试多种DOC文件提取方法
            extracted = False
            
            # 方法1：尝试使用python-docx处理DOC文件（有时也能处理.doc）
            try:
                doc = Document(temp_file_path)
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        lines.append(text)
                        html_content.append(f"<p>{text}</p>")
                for table in doc.tables:
                    html_table = '<table border="1" class="table">'
                    for row in table.rows:
                        html_table += '<tr>'
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            lines.append(cell_text)
                            html_table += f'<td>{cell_text}</td>'
                        html_table += '</tr>'
                    html_table += '</table>'
                    html_content.append(html_table)
                extracted = True
                print("使用python-docx成功处理DOC文件")
            except Exception as e:
                print(f"python-docx处理DOC失败: {e}")
            
            # 方法2：使用textract（如果可用）
            if not extracted:
                try:
                    import textract
                    text = textract.process(temp_file_path).decode('utf-8')
                    cleaned_text = clean_text(text)
                    lines = cleaned_text.splitlines()
                    html_content = [f"<pre>{line}</pre>" for line in lines]
                    extracted = True
                    print("使用textract成功处理DOC文件")
                except ImportError:
                    print("textract不可用")
                except Exception as e:
                    print(f"textract处理DOC失败: {e}")
            
            # 方法3：使用docx2txt（如果可用）
            if not extracted:
                try:
                    import docx2txt
                    text = docx2txt.process(temp_file_path)
                    if text:
                        cleaned_text = clean_text(text)
                        lines = cleaned_text.splitlines()
                        html_content = [f"<pre>{line}</pre>" for line in lines]
                        extracted = True
                        print("使用docx2txt成功处理DOC文件")
                except ImportError:
                    print("docx2txt不可用")
                except Exception as e:
                    print(f"docx2txt处理DOC失败: {e}")

            # 方法4：使用antiword（Linux常用）
            if not extracted:
                try:
                    import subprocess
                    result = subprocess.run(['antiword', temp_file_path], capture_output=True, text=True, check=True)
                    text = result.stdout
                    if text:
                        cleaned_text = clean_text(text)
                        lines = cleaned_text.splitlines()
                        html_content = [f"<pre>{line}</pre>" for line in lines]
                        extracted = True
                        print("使用antiword成功处理DOC文件")
                except Exception as e:
                    print(f"antiword处理DOC失败: {e}")

            # 方法5：使用catdoc兜底
            if not extracted:
                try:
                    import subprocess
                    result = subprocess.run(['catdoc', '-w', temp_file_path], capture_output=True, text=True, check=True)
                    text = result.stdout
                    if text:
                        cleaned_text = clean_text(text)
                        lines = cleaned_text.splitlines()
                        html_content = [f"<pre>{line}</pre>" for line in lines]
                        extracted = True
                        print("使用catdoc成功处理DOC文件")
                except Exception as e:
                    print(f"catdoc处理DOC失败: {e}")
            
            if not extracted:
                print("所有DOC提取方法都失败")
                return None, None, None

        elif ext == 'txt':
            with open(temp_file_path, 'r', encoding='utf-8') as f:
                text = f.read()
                cleaned_text = clean_text(text)
                lines = cleaned_text.splitlines()
                html_content = [f"<pre>{line}</pre>" for line in lines]

        else:
            return None, None, None

        return lines, html_content, temp_file_path

    except Exception as e:
        print(f"提取文本失败: {e}")
        return None, None, None

def segment_text(text, max_sentences=MAX_SENTENCES_PER_SEGMENT):
    """将文本分割为段落"""
    try:
        sentences = _sent_tokenize(text)
    except Exception:
        sentences = _simple_sent_tokenize(text)
    segments = []
    for i in range(0, len(sentences), max_sentences):
        segment = ' '.join(sentences[i:i + max_sentences])
        if segment.strip():
            segments.append(segment)
    return segments if segments else [text]

# 相似度计算：优先ST模型，失败则TF-IDF
try:
    from sentence_transformers import util as _st_util
except ImportError:
    # 如果无法导入，创建一个简单的余弦相似度函数
    import torch
    import torch.nn.functional as F
    class SimpleUtil:
        @staticmethod
        def cos_sim(a, b):
            if isinstance(a, torch.Tensor) and isinstance(b, torch.Tensor):
                # 处理不同维度的张量
                if a.dim() == 1 and b.dim() == 1:
                    # 两个一维张量，直接计算余弦相似度
                    return F.cosine_similarity(a.unsqueeze(0), b.unsqueeze(0), dim=-1)
                elif a.dim() == 1 and b.dim() == 2:
                    # a是一维，b是二维，扩展a的维度
                    return F.cosine_similarity(a.unsqueeze(0), b, dim=-1)
                elif a.dim() == 2 and b.dim() == 1:
                    # a是二维，b是一维，扩展b的维度
                    return F.cosine_similarity(a, b.unsqueeze(0), dim=-1)
                elif a.dim() == 2 and b.dim() == 2:
                    # 两个二维张量，计算矩阵相似度
                    # 使用矩阵乘法计算余弦相似度
                    a_norm = F.normalize(a, p=2, dim=1)
                    b_norm = F.normalize(b, p=2, dim=1)
                    return torch.mm(a_norm, b_norm.t())
                else:
                    # 其他情况，尝试默认处理
                    return F.cosine_similarity(a, b, dim=-1)
            else:
                # 如果不是tensor，尝试转换
                if not isinstance(a, torch.Tensor):
                    a = torch.tensor(a, dtype=torch.float32)
                if not isinstance(b, torch.Tensor):
                    b = torch.tensor(b, dtype=torch.float32)
                return F.cosine_similarity(a, b, dim=-1)
    _st_util = SimpleUtil()

def _similarity_st(segments1, segments2):
    sims = []
    with torch.no_grad():
        for seg1 in segments1:
            for seg2 in segments2:
                e1 = _model.encode(seg1, convert_to_tensor=True)
                e2 = _model.encode(seg2, convert_to_tensor=True)
                sim = _st_util.cos_sim(e1, e2)[0][0].item()
                sims.append(sim)
    return sims

def _similarity_tfidf(segments1, segments2):
    if _tfidf is None:
        return [0.0]
    sims = []
    for seg1 in segments1:
        for seg2 in segments2:
            mats = _tfidf.fit_transform([seg1, seg2])
            sim = float(_sk_cos_sim(mats[0], mats[1])[0, 0])
            sims.append(sim)
    return sims


def compute_similarity(text1, text2):
    """计算两个文本的相似度（分段处理，结合差异分析）。
    优化：一次性批量编码并矩阵相似度，充分利用BLAS/OMP多核或GPU。"""
    if not text1 or not text2:
        return 0.0
    
    # 预处理：规范化文本
    text1_norm = normalize_whitespace(text1)
    text2_norm = normalize_whitespace(text2)
    
    # 如果文本完全相同，直接返回1.0
    if text1_norm == text2_norm:
        return 1.0
    
    # 快速字符级差异分析（避免循环依赖）
    import difflib
    total_orig_chars = len(text1_norm)
    total_new_chars = len(text2_norm)
    total_chars = max(total_orig_chars, total_new_chars, 1)
    
    print(f"🔢 相似度计算: 原文{total_orig_chars}字符, 新文{total_new_chars}字符")
    
    # 使用difflib进行简单的字符级差异统计
    matcher = difflib.SequenceMatcher(None, text1_norm, text2_norm)
    similarity_ratio = matcher.ratio()
    
    # 基于字符差异的相似度
    diff_based_similarity = similarity_ratio
    print(f"🔢 字符级相似度: {diff_based_similarity:.4f}")
    
    # 计算行级相似度（对表格文档更准确）
    lines1 = [line.strip() for line in text1_norm.split('\n') if line.strip()]
    lines2 = [line.strip() for line in text2_norm.split('\n') if line.strip()]
    
    line_matcher = difflib.SequenceMatcher(None, lines1, lines2)
    line_based_similarity = line_matcher.ratio()
    print(f"🔢 行级相似度: {line_based_similarity:.4f}")
    
    # 对于有明显差异的文档，使用更严格的相似度
    if diff_based_similarity < 0.98 or line_based_similarity < 0.98:
        # 取更低的相似度值，确保敏感性
        strict_similarity = min(diff_based_similarity, line_based_similarity)
        
        # 长度差异惩罚
        length_diff = abs(total_orig_chars - total_new_chars)
        if length_diff > 0:
            length_penalty = 1.0 - (length_diff / total_chars * 0.5)
            strict_similarity *= max(0.7, length_penalty)
            print(f"🔢 长度差异惩罚: {length_penalty:.4f}")
        
        print(f"🔢 严格相似度（提前返回）: {strict_similarity:.4f}")
        return max(0.0, min(1.0, strict_similarity))
    
    # 传统的段落相似度计算
    segments1 = segment_text(text1)
    segments2 = segment_text(text2)
    
    # 显示计算模式信息（仅在调试时）
    if os.environ.get('DEBUG_SIMILARITY', '0') == '1':
        if _model is not None and hasattr(_model, 'device'):
            device_info = getattr(_model, 'device', 'unknown')
            if str(device_info).startswith('cuda'):
                print(f"🚀 GPU加速计算相似度 - 设备: {device_info}")
            else:
                print(f"🖥️  CPU多线程计算相似度 - 段落数: {len(segments1)}x{len(segments2)}")
        elif _use_tfidf_fallback:
            print(f"📊 TF-IDF计算相似度 - 段落数: {len(segments1)}x{len(segments2)}")

    if _model is not None:
        # SentenceTransformer 路径（GPU或CPU）
        try:
            with torch.no_grad():
                # 激进的内存管理策略
                if torch.cuda.is_available():
                    # 检查GPU内存使用情况
                    gpu_memory_used = torch.cuda.memory_allocated() / 1024**3  # GB
                    gpu_memory_total = torch.cuda.get_device_properties(0).total_memory / 1024**3  # GB
                    gpu_memory_free = gpu_memory_total - gpu_memory_used
                    
                    # 根据可用内存动态调整批次大小
                    if gpu_memory_free < 1.0:  # 小于1GB可用内存
                        max_batch_size = 4
                    elif gpu_memory_free < 2.0:  # 小于2GB可用内存
                        max_batch_size = 8
                    elif gpu_memory_free < 3.0:  # 小于3GB可用内存
                        max_batch_size = 16
                    else:
                        max_batch_size = 32
                else:
                    max_batch_size = 64
                
                batch_size = min(max_batch_size, max(len(segments1), len(segments2)))
                
                # 如果文本太长，进一步减少批次大小
                if len(segments1) > 30 or len(segments2) > 30:
                    batch_size = min(batch_size, 8)
                if len(segments1) > 100 or len(segments2) > 100:
                    batch_size = min(batch_size, 4)
                
                # 清理GPU缓存
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
                
                # 检查模型是否有device属性，如果没有则使用默认设备
                model_device = getattr(_model, 'device', device)
                emb1 = _model.encode(segments1, convert_to_tensor=True, show_progress_bar=False, 
                                   batch_size=batch_size)
                emb2 = _model.encode(segments2, convert_to_tensor=True, show_progress_bar=False, 
                                   batch_size=batch_size)
                
                # 确保张量在正确的设备上
                if hasattr(emb1, 'to'):
                    emb1 = emb1.to(model_device)
                if hasattr(emb2, 'to'):
                    emb2 = emb2.to(model_device)
                # 余弦相似度矩阵 (len1 x len2)
                sim_mat = _st_util.cos_sim(emb1, emb2)
                max_sim = float(sim_mat.max().item()) if sim_mat.numel() > 0 else 0.0
                
                # 立即清理GPU缓存
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
        except Exception as e:
            print(f"ST批量相似度计算失败，回退到逐对: {e}")
            sims = []
            with torch.no_grad():
                for s1 in segments1:
                    try:
                        e1 = _model.encode(s1, convert_to_tensor=True, batch_size=1)
                        for s2 in segments2:
                            try:
                                e2 = _model.encode(s2, convert_to_tensor=True, batch_size=1)
                                # 确保两个张量都是2D的
                                if e1.dim() == 1:
                                    e1 = e1.unsqueeze(0)
                                if e2.dim() == 1:
                                    e2 = e2.unsqueeze(0)
                                sim = float(_st_util.cos_sim(e1, e2)[0][0].item())
                                sims.append(sim)
                                # 清理中间变量
                                del e2
                                if torch.cuda.is_available():
                                    torch.cuda.empty_cache()
                            except Exception as e2:
                                print(f"逐对计算失败 (s2): {e2}")
                                continue
                        # 清理中间变量
                        del e1
                        if torch.cuda.is_available():
                            torch.cuda.empty_cache()
                    except Exception as e1:
                        print(f"逐对计算失败 (s1): {e1}")
                        continue
            max_sim = max(sims) if sims else 0.0
    else:
        # TF-IDF 回退路径（批量构造再两两计算）
        if _tfidf is None:
            return 0.0
        try:
            docs = segments1 + segments2
            mats = _tfidf.fit_transform(docs)  # (n1+n2, vocab)
            n1 = len(segments1)
            A = mats[:n1]
            B = mats[n1:]
            # 计算余弦相似度矩阵 A * B^T
            from sklearn.metrics.pairwise import cosine_similarity as _sk_cos_sim
            sim_mat = _sk_cos_sim(A, B)
            max_sim = float(sim_mat.max()) if sim_mat.size > 0 else 0.0
        except Exception as e:
            print(f"TF-IDF批量相似度计算失败: {e}")
            # 回退到逐对
            sims = []
            for s1 in segments1:
                for s2 in segments2:
                    try:
                        mats = _tfidf.fit_transform([s1, s2])
                        sims.append(float((_sk_cos_sim(mats[0], mats[1]))[0, 0]))
                    except Exception:
                        sims.append(0.0)
            max_sim = max(sims) if sims else 0.0

    # 组合相似度：取差异分析相似度和传统相似度的加权平均
    # 对于有明显差异的文档，差异分析更准确
    # 对于内容相近的文档，传统方法更准确
    if abs(diff_based_similarity - 1.0) > 0.01:
        # 有差异时，更依赖差异分析
        final_similarity = 0.6 * diff_based_similarity + 0.4 * max_sim
    else:
        # 无差异时，使用传统方法
        final_similarity = max_sim
    
    # 确保相似度在合理范围内
    final_similarity = max(0.0, min(1.0, final_similarity))
    
    if final_similarity > SIMILARITY_THRESHOLD:
        print(f"警告：相似度 {final_similarity:.4f} 超过阈值 {SIMILARITY_THRESHOLD}")
    
    return final_similarity

def get_common_words(text1, text2, top_n=20):
    """获取两个文本中相同且高频次出现的词汇，排除纯数字"""
    try:
        words1 = [word.lower() for word in _word_tokenize(text1) if word.lower().isalnum() and not word.isdigit()]
        words2 = [word.lower() for word in _word_tokenize(text2) if word.lower().isalnum() and not word.isdigit()]
    except Exception:
        words1 = _simple_word_tokenize(text1)
        words2 = _simple_word_tokenize(text2)
        words1 = [w for w in words1 if w.isalnum() and not w.isdigit()]
        words2 = [w for w in words2 if w.isalnum() and not w.isdigit()]
    common_words = set(words1) & set(words2)
    from collections import Counter as _Counter
    word_counts1 = _Counter(word for word in words1 if word in common_words)
    word_counts2 = _Counter(word for word in words2 if word in common_words)
    combined_counts = {word: min(word_counts1[word], word_counts2[word]) for word in common_words}
    return sorted(combined_counts.items(), key=lambda x: x[1], reverse=True)[:top_n]

def get_similarity_reason(text1, text2, common_words, similarity):
    """动态生成定制化的相似原因"""
    if not common_words:
        return "相似度低，文件内容可能无显著相关性。"
    similarity_percent = similarity * 100
    reason = f"相似度为 {similarity_percent:.2f}%，"
    if similarity_percent > 50:
        reason += "高度相关，主要由以下共享词汇驱动："
    elif similarity_percent < 30:
        reason += "低相关，可能由少量共享词汇引起，内容差异较大："
    else:
        reason += "部分相关，主要由以下共享词汇引起："
    total_words = len(set(word for word, _ in common_words))
    significant_words = [(word, count) for word, count in common_words if count > 2 or (count / total_words) > 0.1]
    if significant_words:
        key_terms = ", ".join(word for word, _ in significant_words[:3])
        reason += f" {key_terms} 等，"
        avg_length = sum(len(word) for word, _ in significant_words) / len(significant_words)
        if avg_length > 5 and any(len(word) > 6 for word, _ in significant_words):
            reason += "可能反映技术或专业领域相关内容。"
        elif len(significant_words) / total_words > 0.3:
            reason += "可能反映特定主题或配置相关内容。"
        else:
            reason += "可能涉及多种主题，相关性需进一步验证。"
    else:
        reason += "共享词汇分布不集中，内容相关性不明显。"
    return reason

def load_similarities():
    """加载现有相似度JSON"""
    if os.path.exists(SIMILARITIES_FILE):
        with open(SIMILARITIES_FILE, 'r') as f:
            return json.load(f)
    return {}

def save_similarities(data):
    """保存相似度JSON"""
    with open(SIMILARITIES_FILE, 'w') as f:
        json.dump(data, f, indent=4)

def smart_split_text(text, max_chunk_size=1000):
    """智能分割文本，优先按段落、句子、行分割"""
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    current_chunk = ""
    lines = text.splitlines()
    
    for line in lines:
        # 如果当前行加上当前块超过最大大小，保存当前块
        if len(current_chunk) + len(line) + 1 > max_chunk_size and current_chunk:
            chunks.append(current_chunk.strip())
            current_chunk = line
        else:
            if current_chunk:
                current_chunk += "\n" + line
            else:
                current_chunk = line
    
    # 添加最后一个块
    if current_chunk:
        chunks.append(current_chunk.strip())
    
    return chunks


def compute_differences(text1, text2):
    """计算两个文本的差异（智能分割，多级别比较）"""
    if not text1 and not text2:
        return []
    
    # 🚨 临时禁用表格检测，使用强制字符级逐行差异
    print("🔥 使用强制字符级逐行差异算法")
    return compute_character_level_line_differences(text1, text2)

def compute_character_level_line_differences(text1, text2):
    """
    字符级逐行差异检测 - 彻底解决大段差异问题
    """
    print("🔥 启动字符级逐行差异检测")
    
    # 🚨 问题发现：文档中的\n被当成了一整行！
    # 解决方案：先按真实行分割，然后进一步分割长行
    
    def split_into_real_lines(text):
        """将文本分割成真实的行，处理文本中以字面形式存在的"\\n"与过长行。
        目标：尽量把段落拆到句子级/条目级，以避免大段差异。
        """
        import re
        # 先把字面"\\n"标准化为真正的换行
        normalized_text = text.replace('\\r\\n', '\n').replace('\\n', '\n')
        # 再按换行切分
        initial_lines = normalized_text.split('\n')
        real_lines = []

        for line in initial_lines:
            curr = line.strip()
            if not curr:
                continue

            # 若该行仍很长，分两级：句子级 -> 短语级
            if len(curr) > 100:
                # 句子级：按中文/英文标点切分，并保留标点到片段尾部
                sentence_parts = re.split(r'(?<=[。！？；;:,，、])\s*', curr)
                for sent in sentence_parts:
                    s = sent.strip()
                    if not s:
                        continue
                    if len(s) > 100:
                        # 短语级：进一步按顿号/逗号/冒号/空白块等切分
                        phrase_parts = re.split(r'[、，,：:；;]|\s{2,}|\t+', s)
                        for ph in phrase_parts:
                            p = ph.strip()
                            if not p:
                                continue
                            # 仍过长则按数字/单位混合模式进一步切（表格式数据友好）
                            if len(p) > 100:
                                number_parts = re.split(r'(?:\d+\.?\d*\s*){1,}', p)
                                for np in number_parts:
                                    nps = np.strip()
                                    if nps:
                                        real_lines.append(nps)
                            else:
                                real_lines.append(p)
                    else:
                        real_lines.append(s)
            else:
                real_lines.append(curr)

        return real_lines

    def split_into_clauses(s: str):
        """把一行切成更小的语义片段（句子/短语/表格单元）。"""
        import re
        if not s:
            return []
        # 规范化字面"\\n"并去两端空白
        s = s.replace('\\r\\n', '\n').replace('\\n', '\n').strip()
        # 先按句子结束符/换行切分（保留标点在片段尾部）
        parts = re.split(r'(?<=[。！？；;])\s*|\n+', s)
        clauses = []
        for part in parts:
            t = part.strip()
            if not t:
                continue
            if len(t) > 80:
                # 对较长片段进一步用逗号/顿号/冒号等细分
                subparts = re.split(r'[，、,:：;；]|\s{2,}|\t+', t)
                for sp in subparts:
                    sp = sp.strip()
                    if sp:
                        clauses.append(sp)
            else:
                clauses.append(t)
        return clauses

    def create_clause_level_diffs(orig_line: str, new_line: str, o_idx: int, n_idx: int):
        """在同一行内，基于子句（句子/短语）生成最小差异块。"""
        from difflib import SequenceMatcher as _Seq
        diffs = []
        orig_clauses = split_into_clauses(orig_line)
        new_clauses = split_into_clauses(new_line)
        matcher_c = _Seq(None, orig_clauses, new_clauses)
        for ctag, oi1, oi2, nj1, nj2 in matcher_c.get_opcodes():
            if ctag == 'equal':
                continue
            original_chunk = ''.join(orig_clauses[oi1:oi2]) if oi2 > oi1 else ''
            new_chunk = ''.join(new_clauses[nj1:nj2]) if nj2 > nj1 else ''
            if not original_chunk and not new_chunk:
                continue
            if ctag == 'delete':
                d_type = 'deleted'
            elif ctag == 'insert':
                d_type = 'added'
            else:
                d_type = 'modified'
            diffs.append({
                'id': str(uuid.uuid4()),
                'type': d_type,
                'original': (original_chunk[:200] + ('...' if len(original_chunk) > 200 else '')) if original_chunk else '',
                'new': (new_chunk[:200] + ('...' if len(new_chunk) > 200 else '')) if new_chunk else '',
                'start_idx_orig': o_idx,
                'end_idx_orig': o_idx + (1 if original_chunk else 0),
                'start_idx_new': n_idx,
                'end_idx_new': n_idx + (1 if new_chunk else 0),
                'line_numbers_orig': [o_idx] if original_chunk else [],
                'line_numbers_new': [n_idx] if new_chunk else []
            })
        return diffs
    
    lines1 = split_into_real_lines(text1)
    lines2 = split_into_real_lines(text2)
    
    print(f"🔥 原文{len(lines1)}行，新文{len(lines2)}行")
    
    differences = []
    
    # 使用difflib进行行级比较
    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, lines1, lines2)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
            
        print(f"🔥 发现{tag}差异: 原文{i1}-{i2}行, 新文{j1}-{j2}行")
        
        # 对于每个差异块，强制创建逐行差异
        if tag == 'delete':
            # 删除：原文有内容，新文没有
            for line_idx in range(i1, i2):
                line_content = lines1[line_idx]
                if line_content.strip():  # 只处理非空行
                    diff = {
                        'id': str(uuid.uuid4()),
                        'type': 'deleted',
                        'original': line_content[:200] + ('...' if len(line_content) > 200 else ''),  # 🔥 限制内容长度
                        'new': '',
                        'start_idx_orig': line_idx,
                        'end_idx_orig': line_idx + 1,
                        'start_idx_new': j1,
                        'end_idx_new': j1,
                        'line_numbers_orig': [line_idx],
                        'line_numbers_new': []
                    }
                    differences.append(diff)
                    print(f"🔥 删除行{line_idx}: {line_content[:50]}...")
                    
        elif tag == 'insert':
            # 插入：新文有内容，原文没有
            for line_idx in range(j1, j2):
                line_content = lines2[line_idx]
                if line_content.strip():  # 只处理非空行
                    diff = {
                        'id': str(uuid.uuid4()),
                        'type': 'added',
                        'original': '',
                        'new': line_content[:200] + ('...' if len(line_content) > 200 else ''),  # 🔥 限制内容长度
                        'start_idx_orig': i1,
                        'end_idx_orig': i1,
                        'start_idx_new': line_idx,
                        'end_idx_new': line_idx + 1,
                        'line_numbers_orig': [],
                        'line_numbers_new': [line_idx]
                    }
                    differences.append(diff)
                    print(f"🔥 新增行{line_idx}: {line_content[:50]}...")
                    
        elif tag == 'replace':
            # 替换：逐行比较内容
            orig_lines = lines1[i1:i2]
            new_lines = lines2[j1:j2]
            max_lines = max(len(orig_lines), len(new_lines))
            
            print(f"🔥 替换块：原文{len(orig_lines)}行 -> 新文{len(new_lines)}行，强制逐行分解")
            
            for idx in range(max_lines):
                orig_line = orig_lines[idx] if idx < len(orig_lines) else ""
                new_line = new_lines[idx] if idx < len(new_lines) else ""

                # 行缺失/新增：保持行级粒度
                if not orig_line and new_line:
                    diff = {
                        'id': str(uuid.uuid4()),
                        'type': 'added',
                        'original': '',
                        'new': new_line[:200] + ('...' if len(new_line) > 200 else ''),
                        'start_idx_orig': i1 + idx if idx < len(orig_lines) else i2,
                        'end_idx_orig': i1 + idx if idx < len(orig_lines) else i2,
                        'start_idx_new': j1 + idx if idx < len(new_lines) else j2,
                        'end_idx_new': j1 + idx + 1 if idx < len(new_lines) else j2,
                        'line_numbers_orig': [],
                        'line_numbers_new': [j1 + idx] if idx < len(new_lines) else []
                    }
                    differences.append(diff)
                    print(f"🔥 added差异 原文行{i1 + idx} vs 新文行{j1 + idx}")
                    continue
                if orig_line and not new_line:
                    diff = {
                        'id': str(uuid.uuid4()),
                        'type': 'deleted',
                        'original': orig_line[:200] + ('...' if len(orig_line) > 200 else ''),
                        'new': '',
                        'start_idx_orig': i1 + idx if idx < len(orig_lines) else i2,
                        'end_idx_orig': i1 + idx + 1 if idx < len(orig_lines) else i2,
                        'start_idx_new': j1 + idx if idx < len(new_lines) else j2,
                        'end_idx_new': j1 + idx if idx < len(new_lines) else j2,
                        'line_numbers_orig': [i1 + idx] if idx < len(orig_lines) else [],
                        'line_numbers_new': []
                    }
                    differences.append(diff)
                    print(f"🔥 deleted差异 原文行{i1 + idx} vs 新文行{j1 + idx}")
                    continue

                # 同行修改：进行子句/短语级细分，最小化差异块
                if orig_line != new_line:
                    clause_diffs = create_clause_level_diffs(orig_line, new_line, i1 + idx, j1 + idx)
                    # 回退：若子句级未产生差异，则退回单行modified
                    if not clause_diffs:
                        differences.append({
                            'id': str(uuid.uuid4()),
                            'type': 'modified',
                            'original': orig_line[:200] + ('...' if len(orig_line) > 200 else ''),
                            'new': new_line[:200] + ('...' if len(new_line) > 200 else ''),
                            'start_idx_orig': i1 + idx,
                            'end_idx_orig': i1 + idx + 1,
                            'start_idx_new': j1 + idx,
                            'end_idx_new': j1 + idx + 1,
                            'line_numbers_orig': [i1 + idx],
                            'line_numbers_new': [j1 + idx]
                        })
                        print(f"🔥 modified差异 原文行{i1 + idx} vs 新文行{j1 + idx}")
                    else:
                        differences.extend(clause_diffs)
                        print(f"🔥 子句级细分完成 原文行{i1 + idx} vs 新文行{j1 + idx} → 生成{len(clause_diffs)}个差异片段")
    
    print(f"🔥 字符级分析完成，生成{len(differences)}个逐行差异")
    return differences
    
    # 原有逻辑（暂时禁用）
    # is_structured_doc = detect_structured_document(text1, text2)
    # is_table_doc = detect_table_document(text1, text2)
    # print(f"🔍 文档类型检测: 结构化={is_structured_doc}, 表格={is_table_doc}")
    # if is_structured_doc:
    #     print("📊 使用结构化文档差异算法")
    #     return compute_structured_document_differences(text1, text2)
    # else:
    #     print("📝 使用普通文本差异算法")
    #     return compute_general_text_differences(text1, text2)

def detect_structured_document(text1, text2):
    """检测是否是结构化文档（如Word文档）"""
    # 检测标准：
    # 1. 有较多的短行（可能是标题、列表项等）
    # 2. 有明显的段落结构
    # 3. 内容相对简洁，不是大段连续文本
    # 4. 包含表格类数据特征
    
    def analyze_text_structure(text):
        lines = text.splitlines()
        if len(lines) < 2:
            return False
            
        short_lines = sum(1 for line in lines if len(line.strip()) < 50 and len(line.strip()) > 0)
        total_lines = len([line for line in lines if line.strip()])
        
        if total_lines == 0:
            return False
            
        short_line_ratio = short_lines / total_lines
        avg_line_length = sum(len(line.strip()) for line in lines if line.strip()) / total_lines
        
        # 检测表格特征
        numeric_lines = sum(1 for line in lines if any(char.isdigit() for char in line) and len(line.strip()) > 0)
        numeric_ratio = numeric_lines / total_lines if total_lines > 0 else 0
        
        # 检测重复模式（表格行）
        pattern_lines = 0
        for line in lines:
            if line.strip() and ('\t' in line or line.count(' ') > 5):
                pattern_lines += 1
        pattern_ratio = pattern_lines / total_lines if total_lines > 0 else 0
        
        # 如果短行比例高且平均行长适中，或者有大量数字和表格模式，可能是结构化文档
        return ((short_line_ratio > 0.3 and 20 < avg_line_length < 100) or 
                (numeric_ratio > 0.4) or 
                (pattern_ratio > 0.3))
    
    return analyze_text_structure(text1) or analyze_text_structure(text2)

def detect_table_document(text1, text2):
    """检测是否是表格密集型文档"""
    import re
    
    def has_table_characteristics(text):
        lines = text.splitlines()
        total_lines = len([line for line in lines if line.strip()])
        if total_lines == 0:
            return False
            
        # 检测数字行比例
        numeric_lines = sum(1 for line in lines if re.search(r'\d+.*\d+', line))
        numeric_ratio = numeric_lines / total_lines
        
        # 检测制表符或多空格分隔
        structured_lines = sum(1 for line in lines if '\t' in line or re.search(r'\s{2,}', line))
        structured_ratio = structured_lines / total_lines
        
        # 检测价格、数量等表格特征
        table_pattern_lines = sum(1 for line in lines if re.search(r'(\d+\.?\d*\s+){2,}', line))
        table_ratio = table_pattern_lines / total_lines
        
        print(f"📋 表格检测: 数字行比例={numeric_ratio:.2f}, 结构行比例={structured_ratio:.2f}, 表格模式比例={table_ratio:.2f}")
        
        # 大幅降低阈值，强制启用表格模式
        result = numeric_ratio > 0.1 or structured_ratio > 0.2 or table_ratio > 0.05 or total_lines > 20
        print(f"📋 表格检测结果: {result} (强制表格模式)")
        return result
    
    return has_table_characteristics(text1) or has_table_characteristics(text2)

def compute_structured_document_differences(text1, text2):
    """计算结构化文档的差异（优化的Word文档比对算法）"""
    # 第一步：预处理，规范化空白字符
    text1_normalized = normalize_whitespace(text1)
    text2_normalized = normalize_whitespace(text2)
    
    # 检测是否是表格密集型文档
    is_table_doc = detect_table_document(text1, text2)
    
    if is_table_doc:
        print("📋 检测到表格文档，使用行级比较算法")
        # 表格文档使用行级比较
        return compute_table_document_differences(text1_normalized, text2_normalized)
    else:
        print("📄 检测到普通结构化文档，使用段落级比较算法")
        # 普通结构化文档使用段落级比较
        return compute_paragraph_document_differences(text1_normalized, text2_normalized)

def compute_table_document_differences(text1, text2):
    """专门处理表格密集型文档的差异 - 强制逐行比较模式"""
    lines1 = text1.splitlines()
    lines2 = text2.splitlines()
    
    print(f"📋 表格文档行数: 原文{len(lines1)}行, 新文{len(lines2)}行")
    print(f"📋 启用强制逐行比较模式")
    
    # 过滤空行，但保留索引映射
    non_empty_lines1 = []
    non_empty_lines2 = []
    line_map1 = {}  # 映射过滤后索引到原始索引
    line_map2 = {}
    
    for i, line in enumerate(lines1):
        if line.strip():
            line_map1[len(non_empty_lines1)] = i
            non_empty_lines1.append(line.strip())
    
    for i, line in enumerate(lines2):
        if line.strip():
            line_map2[len(non_empty_lines2)] = i
            non_empty_lines2.append(line.strip())
    
    print(f"📋 过滤空行后: 原文{len(non_empty_lines1)}行, 新文{len(non_empty_lines2)}行")
    
    # 强制使用逐行比较，不允许任何大块差异
    differences = []
    
    # 直接调用细分函数，跳过初始的大块比较
    print(f"📋 直接进行逐行细分比较")
    differences = split_large_table_difference(
        non_empty_lines1, non_empty_lines2, 
        0, 0, line_map1, line_map2
    )
    
    print(f"📋 最终生成{len(differences)}个差异")
    return differences

def split_large_table_difference(lines1, lines2, offset1, offset2, line_map1, line_map2):
    """将大的表格差异块分解为最小粒度的逐行差异"""
    differences = []
    
    print(f"📋 开始细分差异: 原文{len(lines1)}行, 新文{len(lines2)}行")
    
    # 使用更精细的逐行比较算法
    from difflib import SequenceMatcher
    matcher = SequenceMatcher(None, lines1, lines2)
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
            
        print(f"📋 细分差异: {tag}, 原文{i1}-{i2}, 新文{j1}-{j2}")
        
        # 对于replace类型，需要特殊处理 - 逐行比较找出真正不同的行
        if tag == 'replace':
            print(f"📋 处理replace块: 原文{i1}-{i2} vs 新文{j1}-{j2}")
            # 使用逐行比较找出真正的差异
            lines1_chunk = lines1[i1:i2]
            lines2_chunk = lines2[j1:j2]
            
            # 对这个块内部再次使用SequenceMatcher
            chunk_matcher = SequenceMatcher(None, lines1_chunk, lines2_chunk)
            
            # 如果replace块太大，强制逐行分解，不再使用SequenceMatcher
            if (i2 - i1) > 10 or (j2 - j1) > 10:
                print(f"📋 replace块过大，强制逐行分解")
                max_lines = max(i2 - i1, j2 - j1)
                for idx in range(max_lines):
                    line1_idx = i1 + idx if idx < (i2 - i1) else None
                    line2_idx = j1 + idx if idx < (j2 - j1) else None
                    
                    line1 = lines1[line1_idx] if line1_idx is not None else ""
                    line2 = lines2[line2_idx] if line2_idx is not None else ""
                    
                    if line1 != line2:
                        diff = create_single_line_difference(
                            line1, line2,
                            offset1 + (line1_idx if line1_idx is not None else i2),
                            offset2 + (line2_idx if line2_idx is not None else j2),
                            line_map1, line_map2
                        )
                        if diff:
                            differences.append(diff)
                            print(f"📋 强制单行差异: {diff['type']}, 原文行{diff['start_idx_orig']}, 新文行{diff['start_idx_new']}")
                continue  # 跳过后面的chunk_matcher逻辑
            for chunk_tag, ci1, ci2, cj1, cj2 in chunk_matcher.get_opcodes():
                if chunk_tag == 'equal':
                    continue
                    
                print(f"📋 块内差异: {chunk_tag}, 子块{ci1}-{ci2} vs {cj1}-{cj2}")
                
                # 为每个子块创建单行差异
                max_sub_lines = max(ci2 - ci1, cj2 - cj1)
                for sub_idx in range(max_sub_lines):
                    sub_line1_idx = ci1 + sub_idx if sub_idx < (ci2 - ci1) else None
                    sub_line2_idx = cj1 + sub_idx if sub_idx < (cj2 - cj1) else None
                    
                    sub_line1 = lines1_chunk[sub_line1_idx] if sub_line1_idx is not None else ""
                    sub_line2 = lines2_chunk[sub_line2_idx] if sub_line2_idx is not None else ""
                    
                    if sub_line1 != sub_line2:
                        global_line1_idx = i1 + (sub_line1_idx if sub_line1_idx is not None else (ci2 - ci1))
                        global_line2_idx = j1 + (sub_line2_idx if sub_line2_idx is not None else (cj2 - cj1))
                        
                        diff = create_single_line_difference(
                            sub_line1, sub_line2,
                            offset1 + global_line1_idx,
                            offset2 + global_line2_idx,
                            line_map1, line_map2
                        )
                        if diff:
                            differences.append(diff)
                            print(f"📋 创建单行差异: {diff['type']}, 原文行{diff['start_idx_orig']}, 新文行{diff['start_idx_new']}")
        else:
            # 对于delete和insert，直接逐行创建差异
            max_lines = max(i2 - i1, j2 - j1)
            for idx in range(max_lines):
                line1_idx = i1 + idx if idx < (i2 - i1) else None
                line2_idx = j1 + idx if idx < (j2 - j1) else None
                
                line1 = lines1[line1_idx] if line1_idx is not None else ""
                line2 = lines2[line2_idx] if line2_idx is not None else ""
                
                # 只有当行不同时才创建差异
                if line1 != line2:
                    diff = create_single_line_difference(
                        line1, line2, 
                        offset1 + (line1_idx if line1_idx is not None else i2),
                        offset2 + (line2_idx if line2_idx is not None else j2),
                        line_map1, line_map2
                    )
                    if diff:
                        differences.append(diff)
                        print(f"📋 创建单行差异: {diff['type']}, 原文行{diff['start_idx_orig']}, 新文行{diff['start_idx_new']}")
    
    print(f"📋 细分完成，生成{len(differences)}个精细差异")
    return differences

def create_single_line_difference(line1, line2, orig_offset, new_offset, line_map1, line_map2):
    """创建单行差异"""
    diff_id = str(uuid.uuid4())
    
    # 确定差异类型
    if line1 and line2:
        diff_type = 'modified'
    elif line2 and not line1:
        diff_type = 'added'
    elif line1 and not line2:
        diff_type = 'deleted'
    else:
        return None
    
    orig_start = line_map1.get(orig_offset, orig_offset)
    orig_end = orig_start + 1
    
    new_start = line_map2.get(new_offset, new_offset)
    new_end = new_start + 1
    
    return {
        'id': diff_id,
        'type': diff_type,
        'original': line1,
        'new': line2,
        'start_idx_orig': orig_start,
        'end_idx_orig': orig_end,
        'start_idx_new': new_start,
        'end_idx_new': new_end,
        'char_diff': None,
        'line_numbers_orig': [orig_start] if line1 else [],
        'line_numbers_new': [new_start] if line2 else [],
        'context_before_orig': [],
        'context_after_orig': [],
        'context_before_new': [],
        'context_after_new': []
    }

def create_chunk_difference(tag, chunk1, chunk2, orig_offset, new_offset, line_map1, line_map2):
    """创建小块差异"""
    if not chunk1 and not chunk2:
        return None
        
    diff_id = str(uuid.uuid4())
    
    # 确定差异类型
    if tag == 'replace':
        diff_type = 'modified'
    elif tag == 'insert':
        diff_type = 'added'
    elif tag == 'delete':
        diff_type = 'deleted'
    else:
        diff_type = tag
    
    orig_start = line_map1.get(orig_offset, orig_offset)
    orig_end = line_map1.get(orig_offset + len(chunk1) - 1, orig_start) + 1 if chunk1 else orig_start
    
    new_start = line_map2.get(new_offset, new_offset)
    new_end = line_map2.get(new_offset + len(chunk2) - 1, new_start) + 1 if chunk2 else new_start
    
    return {
        'id': diff_id,
        'type': diff_type,
        'original': '\n'.join(chunk1),
        'new': '\n'.join(chunk2),
        'start_idx_orig': orig_start,
        'end_idx_orig': orig_end,
        'start_idx_new': new_start,
        'end_idx_new': new_end,
        'char_diff': None,
        'line_numbers_orig': list(range(orig_start, orig_end)),
        'line_numbers_new': list(range(new_start, new_end)),
        'context_before_orig': [],
        'context_after_orig': [],
        'context_before_new': [],
        'context_after_new': []
    }

def create_table_difference(tag, lines1, lines2, i1, i2, j1, j2, line_map1, line_map2):
    """创建表格差异对象"""
    diff_id = str(uuid.uuid4())
    
    orig_content = '\n'.join(lines1[i1:i2]) if i1 < i2 else ''
    new_content = '\n'.join(lines2[j1:j2]) if j1 < j2 else ''
    
    # 确定差异类型
    if tag == 'replace':
        diff_type = 'modified'
    elif tag == 'insert':
        diff_type = 'added'
    elif tag == 'delete':
        diff_type = 'deleted'
    else:
        diff_type = tag
    
    # 计算原始行号
    orig_start = line_map1.get(i1, 0)
    orig_end = line_map1.get(i2-1, orig_start) + 1 if i2 > 0 and (i2-1) in line_map1 else orig_start + 1
    
    new_start = line_map2.get(j1, 0)
    new_end = line_map2.get(j2-1, new_start) + 1 if j2 > 0 and (j2-1) in line_map2 else new_start + 1
    
    return {
        'id': diff_id,
        'type': diff_type,
        'original': orig_content,
        'new': new_content,
        'start_idx_orig': orig_start,
        'end_idx_orig': orig_end,
        'start_idx_new': new_start,
        'end_idx_new': new_end,
        'char_diff': None,
        'line_numbers_orig': list(range(orig_start, orig_end)),
        'line_numbers_new': list(range(new_start, new_end)),
        'context_before_orig': lines1[max(0, i1-2):i1] if i1 > 0 else [],
        'context_after_orig': lines1[i2:i2+2] if i2 < len(lines1) else [],
        'context_before_new': lines2[max(0, j1-2):j1] if j1 > 0 else [],
        'context_after_new': lines2[j2:j2+2] if j2 < len(lines2) else []
    }

def compute_paragraph_document_differences(text1, text2):
    """处理普通结构化文档的差异"""
    # 第二步：按段落分割
    paragraphs1 = split_into_paragraphs(text1)
    paragraphs2 = split_into_paragraphs(text2)
    
    # 第三步：使用改进的序列匹配算法
    matcher = difflib.SequenceMatcher(None, paragraphs1, paragraphs2)
    differences = []
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        # 对于大的差异块，尝试进一步细分
        if tag == 'replace' and (i2 - i1) > 1 and (j2 - j1) > 1:
            # 在段落内部进一步比较
            sub_diffs = compute_paragraph_internal_differences(
                paragraphs1[i1:i2], paragraphs2[j1:j2], i1, j1
            )
            differences.extend(sub_diffs)
        else:
            # 生成标准差异
            diff = create_paragraph_difference(
                tag, paragraphs1, paragraphs2, i1, i2, j1, j2
            )
            if diff:
                differences.append(diff)
    
    return differences

def compute_general_text_differences(text1, text2):
    """计算普通文本的差异（原有逻辑）"""
    # 优化策略：始终使用多级别分析，选择最佳粒度
    line_diffs = compute_line_level_differences(text1, text2)
    sentence_diffs = compute_sentence_level_differences(text1, text2)
    paragraph_diffs = compute_paragraph_level_differences(text1, text2)
    
    # 计算每种方法的质量分数（差异数量 vs 差异大小的平衡）
    def calculate_quality_score(diffs):
        if not diffs:
            return 0
        
        # 分数 = 差异数量 * 权重 - 平均差异大小的惩罚
        diff_count = len(diffs)
        avg_size = sum(len(d.get('original', '')) + len(d.get('new', '')) for d in diffs) / max(diff_count, 1)
        
        # 差异数量越多越好（更细粒度），但平均大小过大时要惩罚
        quality_score = diff_count * 10 - (avg_size / 100)
        return max(0, quality_score)
    
    line_score = calculate_quality_score(line_diffs)
    sentence_score = calculate_quality_score(sentence_diffs)
    paragraph_score = calculate_quality_score(paragraph_diffs)
    
    # 选择最佳的分割方法
    best_diffs = line_diffs
    best_score = line_score
    
    if sentence_score > best_score:
        best_diffs = sentence_diffs
        best_score = sentence_score
        
    if paragraph_score > best_score:
        best_diffs = paragraph_diffs
        best_score = paragraph_score
    
    # 如果最佳结果仍然不够好（差异过少且过大），尝试智能分割
    if len(best_diffs) <= 2 and best_diffs:
        avg_diff_size = sum(len(d.get('original', '')) + len(d.get('new', '')) for d in best_diffs) / len(best_diffs)
        if avg_diff_size > 800:  # 平均差异太大
            smart_diffs = compute_smart_chunk_differences(text1, text2)
            smart_score = calculate_quality_score(smart_diffs)
            if smart_score > best_score:
                best_diffs = smart_diffs
    
    return best_diffs

def normalize_whitespace(text):
    """规范化空白字符，减少因格式差异导致的误判"""
    import re
    # 统一换行符
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # 移除行首行尾空白，但保留段落间的空行
    lines = text.split('\n')
    normalized_lines = []
    for line in lines:
        if line.strip():
            normalized_lines.append(line.strip())
        else:
            # 保留空行，但只保留一个
            if normalized_lines and normalized_lines[-1] != '':
                normalized_lines.append('')
    
    return '\n'.join(normalized_lines)

def split_into_paragraphs(text):
    """将文本分割成段落，更智能地处理结构化文档"""
    # 按空行分割段落
    paragraphs = []
    current_paragraph = []
    
    lines = text.split('\n')
    for line in lines:
        if line.strip():
            current_paragraph.append(line)
        else:
            # 遇到空行，结束当前段落
            if current_paragraph:
                paragraphs.append('\n'.join(current_paragraph))
                current_paragraph = []
    
    # 添加最后一个段落
    if current_paragraph:
        paragraphs.append('\n'.join(current_paragraph))
    
    return paragraphs

def compute_paragraph_internal_differences(paras1, paras2, offset1, offset2):
    """计算段落内部的精细差异"""
    differences = []
    
    # 将多个段落合并成两个文本块进行比较
    text1 = '\n'.join(paras1)
    text2 = '\n'.join(paras2)
    
    # 使用句子级别比较
    sentences1 = split_into_sentences(text1)
    sentences2 = split_into_sentences(text2)
    
    matcher = difflib.SequenceMatcher(None, sentences1, sentences2)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        diff_id = str(uuid.uuid4())
        
        orig_content = ''.join(sentences1[i1:i2]) if i1 < i2 else ''
        new_content = ''.join(sentences2[j1:j2]) if j1 < j2 else ''
        
        # 确定差异类型
        if tag == 'replace':
            diff_type = 'modified'
        elif tag == 'insert':
            diff_type = 'added'
        elif tag == 'delete':
            diff_type = 'deleted'
        else:
            diff_type = tag
        
        # 计算字符级别的差异
        char_diff = None
        if diff_type == 'modified' and orig_content and new_content:
            char_diff = compute_char_level_diff(orig_content, new_content)
        
        diff = {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': offset1 + i1,
            'end_idx_orig': offset1 + i2,
            'start_idx_new': offset2 + j1,
            'end_idx_new': offset2 + j2,
            'char_diff': char_diff,
            'line_numbers_orig': [],
            'line_numbers_new': [],
            'context_before_orig': sentences1[max(0, i1-1):i1] if i1 > 0 else [],
            'context_after_orig': sentences1[i2:i2+1] if i2 < len(sentences1) else [],
            'context_before_new': sentences2[max(0, j1-1):j1] if j1 > 0 else [],
            'context_after_new': sentences2[j2:j2+1] if j2 < len(sentences2) else []
        }
        
        differences.append(diff)
    
    return differences

def create_paragraph_difference(tag, paras1, paras2, i1, i2, j1, j2):
    """创建段落级别的差异对象"""
    diff_id = str(uuid.uuid4())
    
    orig_content = '\n'.join(paras1[i1:i2]) if i1 < i2 else ''
    new_content = '\n'.join(paras2[j1:j2]) if j1 < j2 else ''
    
    # 确定差异类型
    if tag == 'replace':
        diff_type = 'modified'
    elif tag == 'insert':
        diff_type = 'added'
    elif tag == 'delete':
        diff_type = 'deleted'
    else:
        diff_type = tag
    
    # 计算字符级别的差异
    char_diff = None
    if diff_type == 'modified' and orig_content and new_content:
        char_diff = compute_char_level_diff(orig_content, new_content)
    
    return {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': i1,
            'end_idx_orig': i2,
            'start_idx_new': j1,
            'end_idx_new': j2,
            'char_diff': char_diff,
        'line_numbers_orig': list(range(i1, i2)),
        'line_numbers_new': list(range(j1, j2)),
        'context_before_orig': paras1[max(0, i1-1):i1] if i1 > 0 else [],
        'context_after_orig': paras1[i2:i2+1] if i2 < len(paras1) else [],
        'context_before_new': paras2[max(0, j1-1):j1] if j1 > 0 else [],
        'context_after_new': paras2[j2:j2+1] if j2 < len(paras2) else []
    }


def compute_sentence_level_differences(text1, text2):
    """计算句子级别的差异（优化的分段比对）"""
    # 使用改进的句子分割，更好地处理中文
    sentences1 = split_into_sentences(text1)
    sentences2 = split_into_sentences(text2)
    
    # 如果句子太少，返回空结果让其他方法处理
    if len(sentences1) <= 1 and len(sentences2) <= 1:
        return []
    
    matcher = difflib.SequenceMatcher(None, sentences1, sentences2)
    differences = []
    
    # 用于计算在原始文本中的行号位置
    def find_line_positions(sentences, original_text):
        line_positions = []
        lines = original_text.splitlines()
        current_line = 0
        
        for sentence in sentences:
            sentence_start_line = current_line
            # 在原始行中查找这个句子
            found = False
            for i in range(current_line, len(lines)):
                if sentence.strip() in lines[i]:
                    sentence_start_line = i
                    current_line = i + 1
                    found = True
                    break
            
            if not found:
                # 如果没找到，继续使用当前行号
                sentence_start_line = current_line
                current_line += 1
                
            line_positions.append(sentence_start_line)
        
        return line_positions
    
    line_pos1 = find_line_positions(sentences1, text1)
    line_pos2 = find_line_positions(sentences2, text2)

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        # 生成唯一ID
        diff_id = str(uuid.uuid4())
        
        # 获取差异内容
        orig_content = ''.join(sentences1[i1:i2]) if i1 < i2 else ''
        new_content = ''.join(sentences2[j1:j2]) if j1 < j2 else ''
        
        # 确定差异类型
        if tag == 'replace':
            diff_type = 'modified'
        elif tag == 'insert':
            diff_type = 'added'
        elif tag == 'delete':
            diff_type = 'deleted'
        else:
            diff_type = tag
        
        # 计算在原文中的位置
        start_line_orig = line_pos1[i1] if i1 < len(line_pos1) else len(text1.splitlines())
        end_line_orig = line_pos1[i2-1] + 1 if i2 > 0 and (i2-1) < len(line_pos1) else start_line_orig + 1
        
        start_line_new = line_pos2[j1] if j1 < len(line_pos2) else len(text2.splitlines())
        end_line_new = line_pos2[j2-1] + 1 if j2 > 0 and (j2-1) < len(line_pos2) else start_line_new + 1
        
        # 计算字符级别的差异
        char_diff = None
        if diff_type == 'modified' and orig_content and new_content:
            char_diff = compute_char_level_diff(orig_content, new_content)
        
        diff = {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': start_line_orig,
            'end_idx_orig': end_line_orig,
            'start_idx_new': start_line_new,
            'end_idx_new': end_line_new,
            'char_diff': char_diff,
            'line_numbers_orig': list(range(start_line_orig, end_line_orig)),
            'line_numbers_new': list(range(start_line_new, end_line_new)),
            'context_before_orig': sentences1[max(0, i1-1):i1] if i1 > 0 else [],
            'context_after_orig': sentences1[i2:i2+1] if i2 < len(sentences1) else [],
            'context_before_new': sentences2[max(0, j1-1):j1] if j1 > 0 else [],
            'context_after_new': sentences2[j2:j2+1] if j2 < len(sentences2) else []
        }
        
        differences.append(diff)

    return differences


def split_into_sentences(text):
    """将文本分割成句子（优化的中英文混合处理）"""
    import re
    
    # 改进的句子分割正则表达式，更好地处理中文标点
    # 处理中文句号（。）、问号（？）、感叹号（！）以及英文的句号、问号、感叹号
    sentence_pattern = r'([。！？.!?]+)'
    
    # 分割文本，保留分隔符
    parts = re.split(sentence_pattern, text)
    
    # 重新组合句子和标点符号
    sentences = []
    current_sentence = ""
    
    for i, part in enumerate(parts):
        if part.strip():
            if re.match(sentence_pattern, part):
                # 这是标点符号，添加到当前句子
                current_sentence += part
                if current_sentence.strip():
                    sentences.append(current_sentence.strip())
                current_sentence = ""
            else:
                # 这是文本内容
                current_sentence += part
    
    # 处理最后一个句子（如果没有标点结尾）
    if current_sentence.strip():
        sentences.append(current_sentence.strip())
    
    # 过滤掉空句子和只有标点的句子
    result = []
    for sentence in sentences:
        cleaned = re.sub(r'^[。！？.!?\s]+$', '', sentence)
        if cleaned and len(cleaned.strip()) > 1:
                result.append(sentence)
    
    # 如果分句结果太少，尝试按换行符分割
    if len(result) <= 1 and '\n' in text:
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) > len(result):
            result = lines
    
    return result


def compute_line_level_differences(text1, text2):
    """计算行级别的差异"""
    lines1 = text1.splitlines()
    lines2 = text2.splitlines()
    matcher = difflib.SequenceMatcher(None, lines1, lines2)
    differences = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        # 生成唯一ID
        diff_id = str(uuid.uuid4())
        
        # 获取差异内容
        orig_content = '\n'.join(lines1[i1:i2]) if i1 < i2 else ''
        new_content = '\n'.join(lines2[j1:j2]) if j1 < j2 else ''
        
        # 确定差异类型
        if tag == 'replace':
            diff_type = 'modified'
        elif tag == 'insert':
            diff_type = 'added'
        elif tag == 'delete':
            diff_type = 'deleted'
        else:
            diff_type = tag
        
        # 计算字符级别的差异（用于更精确的高亮）
        char_diff = None
        if diff_type == 'modified' and orig_content and new_content:
            char_diff = compute_char_level_diff(orig_content, new_content)
        
        diff = {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': i1,
            'end_idx_orig': i2,
            'start_idx_new': j1,
            'end_idx_new': j2,
            'char_diff': char_diff,
            'line_numbers_orig': list(range(i1 + 1, i2 + 1)) if i1 < i2 else [],
            'line_numbers_new': list(range(j1 + 1, j2 + 1)) if j1 < j2 else [],
            'context_before_orig': lines1[max(0, i1-2):i1] if i1 > 0 else [],
            'context_after_orig': lines1[i2:i2+2] if i2 < len(lines1) else [],
            'context_before_new': lines2[max(0, j1-2):j1] if j1 > 0 else [],
            'context_after_new': lines2[j2:j2+2] if j2 < len(lines2) else []
        }
        
        differences.append(diff)

    return differences


def compute_paragraph_level_differences(text1, text2):
    """计算段落级别的差异"""
    # 按双换行符分割段落
    paragraphs1 = [p.strip() for p in text1.split('\n\n') if p.strip()]
    paragraphs2 = [p.strip() for p in text2.split('\n\n') if p.strip()]
    
    matcher = difflib.SequenceMatcher(None, paragraphs1, paragraphs2)
    differences = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        # 生成唯一ID
        diff_id = str(uuid.uuid4())
        
        # 获取差异内容
        orig_content = '\n\n'.join(paragraphs1[i1:i2]) if i1 < i2 else ''
        new_content = '\n\n'.join(paragraphs2[j1:j2]) if j1 < j2 else ''
        
        # 确定差异类型
        if tag == 'replace':
            diff_type = 'modified'
        elif tag == 'insert':
            diff_type = 'added'
        elif tag == 'delete':
            diff_type = 'deleted'
        else:
            diff_type = tag
        
        # 计算字符级别的差异
        char_diff = None
        if diff_type == 'modified' and orig_content and new_content:
            char_diff = compute_char_level_diff(orig_content, new_content)
        
        diff = {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': i1,
            'end_idx_orig': i2,
            'start_idx_new': j1,
            'end_idx_new': j2,
            'char_diff': char_diff,
            'line_numbers_orig': [],  # 段落级别不提供行号
            'line_numbers_new': [],
            'context_before_orig': paragraphs1[max(0, i1-1):i1] if i1 > 0 else [],
            'context_after_orig': paragraphs1[i2:i2+1] if i2 < len(paragraphs1) else [],
            'context_before_new': paragraphs2[max(0, j1-1):j1] if j1 > 0 else [],
            'context_after_new': paragraphs2[j2:j2+1] if j2 < len(paragraphs2) else []
        }
        
        differences.append(diff)

    return differences


def compute_smart_chunk_differences(text1, text2):
    """使用智能分割计算差异"""
    # 智能分割文本
    chunks1 = smart_split_text(text1, max_chunk_size=500)
    chunks2 = smart_split_text(text2, max_chunk_size=500)
    
    matcher = difflib.SequenceMatcher(None, chunks1, chunks2)
    differences = []

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        
        # 生成唯一ID
        diff_id = str(uuid.uuid4())
        
        # 获取差异内容
        orig_content = '\n---\n'.join(chunks1[i1:i2]) if i1 < i2 else ''
        new_content = '\n---\n'.join(chunks2[j1:j2]) if j1 < j2 else ''
        
        # 确定差异类型
        if tag == 'replace':
            diff_type = 'modified'
        elif tag == 'insert':
            diff_type = 'added'
        elif tag == 'delete':
            diff_type = 'deleted'
        else:
            diff_type = tag
        
        # 计算字符级别的差异
        char_diff = None
        if diff_type == 'modified' and orig_content and new_content:
            char_diff = compute_char_level_diff(orig_content, new_content)
        
        diff = {
            'id': diff_id,
            'type': diff_type,
            'original': orig_content,
            'new': new_content,
            'start_idx_orig': i1,
            'end_idx_orig': i2,
            'start_idx_new': j1,
            'end_idx_new': j2,
            'char_diff': char_diff,
            'line_numbers_orig': [],  # 智能分割不提供行号
            'line_numbers_new': [],
            'context_before_orig': chunks1[max(0, i1-1):i1] if i1 > 0 else [],
            'context_after_orig': chunks1[i2:i2+1] if i2 < len(chunks1) else [],
            'context_before_new': chunks2[max(0, j1-1):j1] if j1 > 0 else [],
            'context_after_new': chunks2[j2:j2+1] if j2 < len(chunks2) else []
        }
        
        differences.append(diff)

    return differences


def compute_char_level_diff(text1, text2):
    """计算字符级别的差异，用于更精确的高亮显示"""
    try:
        matcher = difflib.SequenceMatcher(None, text1, text2)
        char_diffs = []
        
        for tag, i1, i2, j1, j2 in matcher.get_opcodes():
            if tag == 'equal':
                continue
            
            char_diff = {
                'type': tag,
                'start_orig': i1,
                'end_orig': i2,
                'start_new': j1,
                'end_new': j2,
                'content_orig': text1[i1:i2] if i1 < i2 else '',
                'content_new': text2[j1:j2] if j1 < j2 else ''
            }
            char_diffs.append(char_diff)
        
        return char_diffs
    except Exception as e:
        print(f"字符级别差异计算失败: {e}")
        return None


def search_keywords_in_text(text, keywords, search_mode='exact'):
    """
    在文本中搜索关键词
    
    Args:
        text: 要搜索的文本
        keywords: 关键词字符串
        search_mode: 搜索模式 ('exact', 'fuzzy', 'regex')
    
    Returns:
        tuple: (是否匹配, 匹配的词汇列表, 匹配位置)
    """
    import re
    
    if not text or not keywords:
        return False, [], []
    
    matches = []
    positions = []
    
    if search_mode == 'exact':
        # 精确匹配
        keyword_lower = keywords.lower()
        text_lower = text.lower()
        if keyword_lower in text_lower:
            matches.append(keywords)
            # 找到所有匹配位置
            start = 0
            while True:
                pos = text_lower.find(keyword_lower, start)
                if pos == -1:
                    break
                positions.append((pos, pos + len(keywords)))
                start = pos + 1
    
    elif search_mode == 'fuzzy':
        # 模糊匹配（包含关键词的单词）
        keyword_words = keywords.lower().split()
        text_lower = text.lower()
        for word in keyword_words:
            if word in text_lower:
                matches.append(word)
                # 找到匹配位置
                start = 0
                while True:
                    pos = text_lower.find(word, start)
                    if pos == -1:
                        break
                    positions.append((pos, pos + len(word)))
                    start = pos + 1
    
    elif search_mode == 'regex':
        # 正则表达式匹配
        try:
            pattern = re.compile(keywords, re.IGNORECASE)
            found_matches = pattern.findall(text)
            matches = list(set(found_matches))
            
            # 找到匹配位置
            for match in pattern.finditer(text):
                positions.append((match.start(), match.end()))
        except re.error as e:
            print(f"正则表达式错误: {e}")
            return False, [], []
    
    return len(matches) > 0, matches, positions


def get_context_around_matches(text, positions, context_size=100):
    """
    获取匹配位置周围的上下文
    
    Args:
        text: 原始文本
        positions: 匹配位置列表 [(start, end), ...]
        context_size: 上下文大小（字符数）
    
    Returns:
        list: 上下文片段列表
    """
    if not positions:
        return []
    
    contexts = []
    for start, end in positions:
        # 计算上下文范围
        context_start = max(0, start - context_size)
        context_end = min(len(text), end + context_size)
        
        # 提取上下文
        context = text[context_start:context_end]
        
        # 标记匹配位置
        relative_start = start - context_start
        relative_end = end - context_start
        
        contexts.append({
            'text': context,
            'match_start': relative_start,
            'match_end': relative_end,
            'original_start': start,
            'original_end': end
        })
    
    return contexts


def calculate_match_score(matches, text_length, match_positions):
    """
    计算匹配分数
    
    Args:
        matches: 匹配的词汇列表
        text_length: 文本长度
        match_positions: 匹配位置列表
    
    Returns:
        dict: 包含各种分数的字典
    """
    if not matches:
        return {
            'match_count': 0,
            'match_ratio': 0.0,
            'density_score': 0.0,
            'coverage_score': 0.0
        }
    
    match_count = len(matches)
    match_ratio = match_count / max(text_length, 1) * 100
    
    # 计算密度分数（匹配位置的平均密度）
    if match_positions:
        total_match_length = sum(end - start for start, end in match_positions)
        density_score = total_match_length / max(text_length, 1) * 100
    else:
        density_score = 0.0
    
    # 计算覆盖分数（匹配位置覆盖的文本比例）
    if match_positions:
        covered_positions = set()
        for start, end in match_positions:
            covered_positions.update(range(start, end))
        coverage_score = len(covered_positions) / max(text_length, 1) * 100
    else:
        coverage_score = 0.0
    
    return {
        'match_count': match_count,
        'match_ratio': round(match_ratio, 2),
        'density_score': round(density_score, 2),
        'coverage_score': round(coverage_score, 2)
    }

# 如果之前替换了print函数，现在恢复它
if '_original_print' in globals():
    print = _original_print

def calculate_similarity_and_diff(text1, text2, task_id=None):
    """计算相似度并生成差异（优先使用diff-match-patch）"""
    similarity = compute_similarity(text1, text2)

    if _DMP_AVAILABLE:
        try:
            # 优先使用 diff-match-patch 进行字符级比较
            dmp = dmp_module.diff_match_patch()
            diffs = dmp.diff_main(text1, text2)
            dmp.diff_cleanupSemantic(diffs)
            
            # 生成美观的HTML高亮差异
            pretty_html = dmp.diff_prettyHtml(diffs)
            
            # 为了安全，将HTML中的换行符替换为<br>，并处理空格
            pretty_html = pretty_html.replace('\\n', '<br>').replace(' &para;', '')
            
            return similarity, pretty_html

        except Exception as e:
            print(f"diff-match-patch 执行失败: {e}, 回退到 difflib")
    
    # 如果 diff-match-patch 不可用或失败，回退到原始的 difflib
    d = difflib.Differ()
    diff = d.compare(text1.splitlines(), text2.splitlines())
    
    diff_html = []
    for line in diff:
        if line.startswith('+ '):
            diff_html.append(f'<span class="diff-added">{line[2:]}</span>')
        elif line.startswith('- '):
            diff_html.append(f'<span class="diff-removed">{line[2:]}</span>')
        elif line.startswith('? '):
            # 忽略difflib的提示行
            continue
        else:
            diff_html.append(f'<span>{line[2:]}</span>')
    
    return similarity, '<br>'.join(diff_html)

def extract_text(file_path):
    """从文件提取纯文本"""
    # ... existing code ...

# [重复的全局变量定义已删除，请参考第1305行附近的统一定义]

def get_model_info():
    """获取当前加载的模型信息"""
    return {
        "startup_mode": _startup_mode,
        "model_source": _model_source,
        "load_time_seconds": round(_model_load_time, 2),
        "is_tfidf": _use_tfidf_fallback,
        "device": str(device)
    }

# 启动流程已移至 startup_manager.py 文件中